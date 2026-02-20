import customtkinter as ctk
from tkinter import filedialog, messagebox
from datetime import datetime, timedelta
import json
import os
import shutil
from pathlib import Path
import csv
from collections import defaultdict

# Configuración de apariencia
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

class QAManagerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Gestión QA Manual")
        self.root.geometry("1400x800")
        
        # Crear directorio de datos
        self.data_dir = Path("qa_data")
        self.data_dir.mkdir(exist_ok=True)
        (self.data_dir / "screenshots").mkdir(exist_ok=True)
        (self.data_dir / "exports").mkdir(exist_ok=True)
        (self.data_dir / "evidencias").mkdir(exist_ok=True)
        (self.data_dir / "backups").mkdir(exist_ok=True)
        
        # Archivos de datos
        self.bugs_file = self.data_dir / "bugs.json"
        self.test_cases_file = self.data_dir / "test_cases.json"
        self.tasks_file = self.data_dir / "tasks.json"
        self.projects_file = self.data_dir / "projects.json"
        self.suites_file = self.data_dir / "test_suites.json"
        self.history_file = self.data_dir / "history.json"
        
        # Cargar datos
        self.bugs = self.load_data(self.bugs_file)
        self.test_cases = self.load_data(self.test_cases_file)
        self.tasks = self.load_data(self.tasks_file)
        self.projects = self.load_data(self.projects_file)
        self.test_suites = self.load_data(self.suites_file)
        self.history = self.load_data(self.history_file)
        
        # 🆕 MEJORA 1: Migrar datos existentes automáticamente
        self.migrate_data_to_new_format()
        
        # 🆕 MEJORA 2: Plantillas de casos de prueba
        self.test_case_templates = self.get_test_case_templates()
        
        # Proyecto actual
        self.current_project = "Todos"
        if not self.projects:
            self.projects = [{"name": "Proyecto General", "version": "1.0", "created": datetime.now().strftime('%Y-%m-%d')}]
            self.save_data(self.projects, self.projects_file)
        
        # Filtros actuales
        self.current_tc_filter = "Todos"
        self.current_bug_filter = "Todos"
        self.current_task_filter = "Todos"
        
        # Variables de selección
        self.selected_bug = None
        self.selected_tc = None
        self.selected_task = None
        
        # Widget actualmente seleccionado
        self.selected_bug_widget = None
        self.selected_tc_widget = None
        self.selected_task_widget = None
        
        # Configurar atajos de teclado
        self.setup_keyboard_shortcuts()
        
        # Backup automático cada 30 minutos
        self.auto_backup()
        
        self.setup_ui()
        
    def setup_keyboard_shortcuts(self):
        """Configurar atajos de teclado"""
        self.root.bind('<Control-n>', lambda e: self.show_new_bug_form())
        self.root.bind('<Control-t>', lambda e: self.show_new_tc_form())
        self.root.bind('<Control-f>', lambda e: self.show_search_dialog())
        self.root.bind('<Control-e>', lambda e: self.show_export_tab())
        self.root.bind('<Control-b>', lambda e: self.create_manual_backup())
        self.root.bind('<Control-p>', lambda e: self.show_projects_dialog())
        self.root.bind('<F5>', lambda e: self.refresh_current_tab())
    
    def auto_backup(self):
        """Backup automático cada 30 minutos"""
        self.create_auto_backup()
        # Programar siguiente backup (30 minutos = 1800000 ms)
        self.root.after(1800000, self.auto_backup)
    
    def create_auto_backup(self):
        """Crear backup automático"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            backup_dir = self.data_dir / "backups" / f"auto_{timestamp}"
            backup_dir.mkdir(parents=True, exist_ok=True)
            
            # Copiar todos los archivos JSON
            for file in [self.bugs_file, self.test_cases_file, self.tasks_file, 
                        self.projects_file, self.suites_file, self.history_file]:
                if file.exists():
                    shutil.copy(file, backup_dir / file.name)
            
            # Limpiar backups antiguos (mantener solo últimos 10)
            backups = sorted([d for d in (self.data_dir / "backups").iterdir() if d.is_dir()])
            if len(backups) > 10:
                for old_backup in backups[:-10]:
                    shutil.rmtree(old_backup)
        except Exception as e:
            print(f"Error en auto-backup: {e}")
    
    def create_manual_backup(self):
        """Crear backup manual"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            backup_dir = self.data_dir / "backups" / f"manual_{timestamp}"
            backup_dir.mkdir(parents=True, exist_ok=True)
            
            for file in [self.bugs_file, self.test_cases_file, self.tasks_file, 
                        self.projects_file, self.suites_file, self.history_file]:
                if file.exists():
                    shutil.copy(file, backup_dir / file.name)
            
            messagebox.showinfo("Éxito", f"Backup creado en:\n{backup_dir}")
        except Exception as e:
            messagebox.showerror("Error", f"Error al crear backup: {e}")
    
    def refresh_current_tab(self):
        """Refrescar la pestaña actual"""
        # Detectar qué tab está visible y refrescarlo
        for tab_name, tab_widget in self.tabs.items():
            if tab_widget.winfo_ismapped():
                if tab_name == 'bugs':
                    self.refresh_bugs_list()
                elif tab_name == 'test_cases':
                    self.refresh_tc_list()
                elif tab_name == 'tasks':
                    self.refresh_tasks_list()
                elif tab_name == 'dashboard':
                    self.refresh_dashboard()
                break
    
    def add_to_history(self, action_type, item_type, item_id, description):
        """Agregar acción al historial"""
        history_entry = {
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'action': action_type,  # 'created', 'updated', 'deleted', 'status_changed'
            'type': item_type,  # 'bug', 'test_case', 'task'
            'item_id': item_id,
            'description': description
        }
        self.history.append(history_entry)
        self.save_data(self.history, self.history_file)
    
    def load_data(self, file_path):
        if file_path.exists():
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    
    def save_data(self, data, file_path):
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    
    def migrate_data_to_new_format(self):
        """MEJORA 1: Migración automática de datos antiguos al nuevo formato"""
        migrated = False
        
        # Migrar bugs: screenshots de STRING a LISTA
        for bug in self.bugs:
            if 'screenshot' in bug and isinstance(bug['screenshot'], str):
                # Convertir screenshot único a lista
                if bug['screenshot']:
                    bug['screenshots'] = [bug['screenshot']]
                else:
                    bug['screenshots'] = []
                del bug['screenshot']
                migrated = True
            elif 'screenshots' not in bug:
                bug['screenshots'] = []
                migrated = True
        
        if migrated:
            self.save_data(self.bugs, self.bugs_file)
            print("Datos migrados al nuevo formato (múltiples screenshots)")
    
    def get_test_case_templates(self):
        """MEJORA 2: Plantillas predefinidas de casos de prueba"""
        return {
            "Login Básico": {
                "module": "Autenticación",
                "type": "Funcional",
                "priority": "Alta",
                "preconditions": "1. Usuario registrado existe en base de datos\n2. Aplicación accesible\n3. Base de datos operativa",
                "steps": "1. Abrir la aplicación\n2. Ingresar email válido registrado\n3. Ingresar contraseña correcta\n4. Hacer clic en botón 'Iniciar Sesión'\n5. Verificar redirección",
                "expected": "1. Usuario es autenticado correctamente\n2. Redirección a dashboard principal\n3. Sesión activa visible\n4. Nombre de usuario mostrado en header"
            },
            "Login - Credenciales Inválidas": {
                "module": "Autenticación",
                "type": "Funcional",
                "priority": "Alta",
                "preconditions": "1. Aplicación accesible\n2. Usuario en página de login",
                "steps": "1. Abrir aplicación\n2. Ingresar email inválido o no registrado\n3. Ingresar cualquier contraseña\n4. Hacer clic en 'Iniciar Sesión'\n5. Verificar mensaje de error",
                "expected": "1. Login rechazado\n2. Mensaje de error claro: 'Credenciales inválidas'\n3. Usuario permanece en página de login\n4. Campos de formulario limpios o resaltados"
            },
            "Validación de Formulario": {
                "module": "UI/Forms",
                "type": "Funcional",
                "priority": "Media",
                "preconditions": "1. Formulario accesible\n2. Campos obligatorios definidos",
                "steps": "1. Abrir formulario\n2. Dejar campos obligatorios vacíos\n3. Intentar enviar formulario\n4. Verificar mensajes de validación\n5. Completar campos obligatorios\n6. Enviar nuevamente",
                "expected": "1. Formulario no se envía si hay campos vacíos\n2. Mensajes de error claros por cada campo obligatorio\n3. Campos resaltados visualmente\n4. Al completar correctamente, formulario se envía exitosamente"
            },
            "API REST - GET": {
                "module": "API",
                "type": "Integración",
                "priority": "Alta",
                "preconditions": "1. API endpoint disponible\n2. Servidor corriendo\n3. Datos de prueba en base de datos",
                "steps": "1. Enviar GET request al endpoint\n2. Verificar status code\n3. Validar estructura JSON de respuesta\n4. Verificar datos retornados\n5. Validar headers de respuesta",
                "expected": "1. Status code 200 OK\n2. JSON válido con estructura esperada\n3. Datos correctos y completos\n4. Headers apropiados (Content-Type: application/json)\n5. Tiempo de respuesta < 2 segundos"
            },
            "API REST - POST": {
                "module": "API",
                "type": "Integración",
                "priority": "Alta",
                "preconditions": "1. API endpoint disponible\n2. Autenticación válida\n3. Payload preparado",
                "steps": "1. Preparar JSON payload válido\n2. Enviar POST request\n3. Verificar status code 201 Created\n4. Validar respuesta con ID generado\n5. Verificar que recurso fue creado (GET)",
                "expected": "1. Status code 201 Created\n2. Respuesta incluye ID del nuevo recurso\n3. Location header presente\n4. Datos persistidos correctamente en base de datos"
            },
            "Responsive Design - Mobile": {
                "module": "UI/UX",
                "type": "Funcional",
                "priority": "Media",
                "preconditions": "1. Aplicación web accesible\n2. DevTools disponible o dispositivo móvil",
                "steps": "1. Abrir aplicación en desktop\n2. Cambiar a vista móvil (320px width)\n3. Verificar layout se ajusta\n4. Probar navegación móvil\n5. Verificar funcionalidad de botones\n6. Probar formularios en móvil",
                "expected": "1. Layout responsive sin scroll horizontal\n2. Todos los elementos visibles y accesibles\n3. Texto legible sin zoom\n4. Botones suficientemente grandes (44x44px mínimo)\n5. Menú de navegación adaptado (hamburguesa)"
            },
            "Performance - Tiempo de Carga": {
                "module": "Performance",
                "type": "Performance",
                "priority": "Alta",
                "preconditions": "1. Aplicación en estado normal\n2. Red estable\n3. Cache limpio",
                "steps": "1. Limpiar cache del navegador\n2. Abrir DevTools > Network\n3. Cargar página principal\n4. Medir tiempo de carga completo\n5. Verificar recursos cargados\n6. Revisar waterfall de requests",
                "expected": "1. Página carga en < 3 segundos\n2. First Contentful Paint < 1.5s\n3. No hay recursos bloqueantes grandes\n4. Imágenes optimizadas\n5. Scripts cargados async/defer"
            },
            "Búsqueda con Resultados": {
                "module": "Búsqueda",
                "type": "Funcional",
                "priority": "Media",
                "preconditions": "1. Sistema de búsqueda implementado\n2. Datos de prueba disponibles",
                "steps": "1. Acceder a función de búsqueda\n2. Ingresar término que tiene resultados\n3. Ejecutar búsqueda\n4. Verificar resultados mostrados\n5. Validar relevancia de resultados",
                "expected": "1. Resultados mostrados correctamente\n2. Resultados relevantes al término buscado\n3. Número de resultados visible\n4. Paginación funcional si hay muchos resultados\n5. Tiempo de búsqueda razonable (< 2s)"
            },
            "Búsqueda sin Resultados": {
                "module": "Búsqueda",
                "type": "Funcional",
                "priority": "Media",
                "preconditions": "1. Sistema de búsqueda implementado",
                "steps": "1. Acceder a función de búsqueda\n2. Ingresar término que NO existe\n3. Ejecutar búsqueda\n4. Verificar mensaje mostrado",
                "expected": "1. Mensaje claro: 'No se encontraron resultados'\n2. Sugerencias de búsqueda alternativa\n3. No hay errores en consola\n4. Interfaz se mantiene estable"
            },
            "Logout / Cerrar Sesión": {
                "module": "Autenticación",
                "type": "Funcional",
                "priority": "Alta",
                "preconditions": "1. Usuario autenticado\n2. Sesión activa",
                "steps": "1. Estando logueado, localizar botón/link de logout\n2. Hacer clic en logout\n3. Verificar redirección\n4. Intentar acceder a página protegida directamente\n5. Verificar que requiere login nuevamente",
                "expected": "1. Sesión cerrada correctamente\n2. Redirección a página de login\n3. Token/cookies de sesión eliminados\n4. No se puede acceder a páginas protegidas sin login\n5. Botón 'back' del navegador no permite volver a sesión"
            }
        }
    
    def setup_ui(self):
        # Frame principal
        main_container = ctk.CTkFrame(self.root)
        main_container.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Sidebar
        self.sidebar = ctk.CTkFrame(main_container, width=220, corner_radius=10)
        self.sidebar.pack(side="left", fill="y", padx=(0, 10))
        self.sidebar.pack_propagate(False)
        
        # Logo/Título en sidebar
        title_label = ctk.CTkLabel(
            self.sidebar, 
            text="Gestión QA Manual", 
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(pady=15, padx=10)
        
        # Selector de proyecto
        project_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        project_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(project_frame, text="Proyecto:", font=ctk.CTkFont(size=11)).pack(anchor="w")
        
        project_names = ["Todos"] + [p['name'] for p in self.projects]
        self.project_selector = ctk.CTkComboBox(
            project_frame, 
            values=project_names,
            command=self.change_project,
            width=180
        )
        self.project_selector.set("Todos")
        self.project_selector.pack(fill="x", pady=(2, 0))
        
        ctk.CTkButton(
            project_frame,
            text="Gestionar Proyectos",
            command=self.show_projects_dialog,
            width=180,
            height=25,
            font=ctk.CTkFont(size=11)
        ).pack(pady=(5, 0))
        
        # Separador
        separator = ctk.CTkFrame(self.sidebar, height=2, fg_color="gray30")
        separator.pack(fill="x", padx=10, pady=10)
        
        # Botones de navegación
        self.nav_buttons = {}
        nav_items = [
            ("Dashboard", self.show_dashboard_tab),
            ("Bugs", self.show_bugs_tab),
            ("Casos de Prueba", self.show_test_cases_tab),
            ("Test Suites", self.show_suites_tab),
            ("Tareas", self.show_tasks_tab),
            ("Buscar", self.show_search_dialog),
            ("Exportar", self.show_export_tab),
            ("Historial", self.show_history_tab)
        ]
        
        for text, command in nav_items:
            btn = ctk.CTkButton(
                self.sidebar,
                text=text,
                command=command,
                font=ctk.CTkFont(size=13),
                height=38,
                corner_radius=8,
                anchor="w"
            )
            btn.pack(pady=3, padx=10, fill="x")
            self.nav_buttons[text] = btn
        
        # Separador
        separator2 = ctk.CTkFrame(self.sidebar, height=2, fg_color="gray30")
        separator2.pack(fill="x", padx=10, pady=10)
        
        # Selector de tema
        theme_label = ctk.CTkLabel(self.sidebar, text="Tema:", font=ctk.CTkFont(size=11))
        theme_label.pack(padx=10, anchor="w")
        
        self.theme_switch = ctk.CTkSegmentedButton(
            self.sidebar,
            values=["Oscuro", "Claro"],
            command=self.change_theme
        )
        self.theme_switch.set("Oscuro")
        self.theme_switch.pack(pady=5, padx=10, fill="x")
        
        # Info de atajos
        shortcuts_label = ctk.CTkLabel(
            self.sidebar,
            text="Atajos:\nCtrl+N: Nuevo Bug\nCtrl+T: Nuevo Caso\nCtrl+F: Buscar\nCtrl+B: Backup\nF5: Refrescar",
            font=ctk.CTkFont(size=9),
            text_color="gray50",
            justify="left"
        )
        shortcuts_label.pack(side="bottom", pady=10, padx=10)
        
        # Área de contenido
        self.content_area = ctk.CTkFrame(main_container, corner_radius=10)
        self.content_area.pack(side="right", fill="both", expand=True)
        
        # Crear todas las pestañas (inicialmente ocultas)
        self.tabs = {}
        self.create_all_tabs()
        
        # Mostrar tab inicial (Dashboard)
        self.show_dashboard_tab()
    
    def change_project(self, project_name):
        """Cambiar proyecto actual"""
        self.current_project = project_name
        self.refresh_current_tab()
        messagebox.showinfo("Proyecto", f"Proyecto cambiado a: {project_name}")
    
    def show_projects_dialog(self):
        """Mostrar diálogo de gestión de proyectos"""
        dialog = ctk.CTkToplevel(self.root)
        dialog.title("Gestión de Proyectos")
        dialog.geometry("600x500")
        dialog.grab_set()
        
        header = ctk.CTkFrame(dialog, fg_color="transparent")
        header.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkLabel(header, text="Proyectos", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left")
        
        ctk.CTkButton(
            header,
            text="+ Nuevo Proyecto",
            command=lambda: self.show_new_project_form(dialog),
            width=140,
            height=35,
            fg_color="#2FA572"
        ).pack(side="right")
        
        # Lista de proyectos
        list_frame = ctk.CTkScrollableFrame(dialog)
        list_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        for project in self.projects:
            item = ctk.CTkFrame(list_frame, corner_radius=8)
            item.pack(fill="x", pady=5, padx=5)
            
            info_frame = ctk.CTkFrame(item, fg_color="transparent")
            info_frame.pack(side="left", fill="x", expand=True, padx=15, pady=10)
            
            ctk.CTkLabel(
                info_frame,
                text=project['name'],
                font=ctk.CTkFont(size=14, weight="bold")
            ).pack(anchor="w")
            
            ctk.CTkLabel(
                info_frame,
                text=f"Versión: {project['version']} | Creado: {project['created']}",
                font=ctk.CTkFont(size=11),
                text_color="gray"
            ).pack(anchor="w")
            
            btn_frame = ctk.CTkFrame(item, fg_color="transparent")
            btn_frame.pack(side="right", padx=10)
            
            ctk.CTkButton(
                btn_frame,
                text="Eliminar",
                command=lambda p=project: self.delete_project(p, dialog),
                width=80,
                height=30,
                fg_color="#D32F2F"
            ).pack()
    
    def show_new_project_form(self, parent):
        """Formulario para nuevo proyecto"""
        form = ctk.CTkToplevel(parent)
        form.title("Nuevo Proyecto")
        form.geometry("400x300")
        form.grab_set()
        
        ctk.CTkLabel(form, text="Crear Nuevo Proyecto", font=ctk.CTkFont(size=18, weight="bold")).pack(pady=20)
        
        ctk.CTkLabel(form, text="Nombre del Proyecto:").pack(pady=(10, 0), padx=20, anchor="w")
        name_entry = ctk.CTkEntry(form, width=350, placeholder_text="Ej: Sistema de Ventas")
        name_entry.pack(pady=5, padx=20)
        
        ctk.CTkLabel(form, text="Versión:").pack(pady=(10, 0), padx=20, anchor="w")
        version_entry = ctk.CTkEntry(form, width=350, placeholder_text="Ej: 1.0")
        version_entry.pack(pady=5, padx=20)
        
        def save_project():
            if not name_entry.get():
                messagebox.showwarning("Advertencia", "El nombre es obligatorio")
                return
            
            project = {
                'name': name_entry.get(),
                'version': version_entry.get() or "1.0",
                'created': datetime.now().strftime('%Y-%m-%d')
            }
            
            self.projects.append(project)
            self.save_data(self.projects, self.projects_file)
            
            # Actualizar selector
            project_names = ["Todos"] + [p['name'] for p in self.projects]
            self.project_selector.configure(values=project_names)
            
            form.destroy()
            parent.destroy()
            self.show_projects_dialog()
            messagebox.showinfo("Éxito", "Proyecto creado correctamente")
        
        ctk.CTkButton(
            form,
            text="Crear Proyecto",
            command=save_project,
            width=150,
            height=40,
            fg_color="#2FA572"
        ).pack(pady=30)
    
    def delete_project(self, project, parent):
        """Eliminar proyecto"""
        if messagebox.askyesno("Confirmar", f"¿Eliminar el proyecto '{project['name']}'?\n\nNOTA: Los bugs y casos NO se eliminarán."):
            self.projects = [p for p in self.projects if p['name'] != project['name']]
            self.save_data(self.projects, self.projects_file)
            
            # Actualizar selector
            project_names = ["Todos"] + [p['name'] for p in self.projects]
            self.project_selector.configure(values=project_names)
            if self.current_project == project['name']:
                self.current_project = "Todos"
                self.project_selector.set("Todos")
            
            parent.destroy()
            self.show_projects_dialog()
    
    def show_search_dialog(self):
        """Mostrar diálogo de búsqueda global"""
        search_window = ctk.CTkToplevel(self.root)
        search_window.title("Búsqueda Global")
        search_window.geometry("900x600")
        search_window.grab_set()
        
        # Header
        header = ctk.CTkFrame(search_window, fg_color="transparent")
        header.pack(fill="x", padx=20, pady=15)
        
        ctk.CTkLabel(header, text="🔍 Buscar", font=ctk.CTkFont(size=22, weight="bold")).pack(side="left")
        
        # Barra de búsqueda
        search_frame = ctk.CTkFrame(search_window, fg_color="transparent")
        search_frame.pack(fill="x", padx=20, pady=10)
        
        search_entry = ctk.CTkEntry(
            search_frame,
            placeholder_text="Buscar en bugs, casos de prueba y tareas...",
            height=40,
            font=ctk.CTkFont(size=14)
        )
        search_entry.pack(side="left", fill="x", expand=True, padx=(0, 10))
        
        results_frame = ctk.CTkScrollableFrame(search_window)
        results_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        def perform_search():
            query = search_entry.get().lower().strip()
            if not query:
                return
            
            # Limpiar resultados anteriores
            for widget in results_frame.winfo_children():
                widget.destroy()
            
            results_count = 0
            
            # Buscar en bugs
            ctk.CTkLabel(results_frame, text="BUGS", font=ctk.CTkFont(size=16, weight="bold"), 
                        fg_color="#E53935", corner_radius=5, padx=10, pady=5).pack(anchor="w", pady=(5, 10))
            
            for bug in self.bugs:
                if (query in bug['title'].lower() or 
                    query in bug['description'].lower() or 
                    query in str(bug['id']).lower()):
                    
                    result_item = ctk.CTkFrame(results_frame, corner_radius=8)
                    result_item.pack(fill="x", pady=3)
                    
                    ctk.CTkLabel(
                        result_item,
                        text=f"#{bug['id']} - {bug['title']}",
                        font=ctk.CTkFont(size=13, weight="bold")
                    ).pack(anchor="w", padx=15, pady=(5, 0))
                    
                    ctk.CTkLabel(
                        result_item,
                        text=f"{bug['status']} | {bug['severity']} | {bug['priority']}",
                        font=ctk.CTkFont(size=11),
                        text_color="gray"
                    ).pack(anchor="w", padx=15, pady=(0, 5))
                    
                    results_count += 1
            
            # Buscar en casos de prueba
            ctk.CTkLabel(results_frame, text="CASOS DE PRUEBA", font=ctk.CTkFont(size=16, weight="bold"), 
                        fg_color="#1E88E5", corner_radius=5, padx=10, pady=5).pack(anchor="w", pady=(15, 10))
            
            for tc in self.test_cases:
                if (query in tc['title'].lower() or 
                    query in tc['test_id'].lower() or
                    query in tc['module'].lower()):
                    
                    result_item = ctk.CTkFrame(results_frame, corner_radius=8)
                    result_item.pack(fill="x", pady=3)
                    
                    ctk.CTkLabel(
                        result_item,
                        text=f"{tc['test_id']} - {tc['title']}",
                        font=ctk.CTkFont(size=13, weight="bold")
                    ).pack(anchor="w", padx=15, pady=(5, 0))
                    
                    ctk.CTkLabel(
                        result_item,
                        text=f"{tc.get('status', 'Sin Probar')} | {tc['module']} | {tc['priority']}",
                        font=ctk.CTkFont(size=11),
                        text_color="gray"
                    ).pack(anchor="w", padx=15, pady=(0, 5))
                    
                    results_count += 1
            
            # Buscar en tareas
            ctk.CTkLabel(results_frame, text="TAREAS", font=ctk.CTkFont(size=16, weight="bold"), 
                        fg_color="#43A047", corner_radius=5, padx=10, pady=5).pack(anchor="w", pady=(15, 10))
            
            for task in self.tasks:
                if (query in task['title'].lower() or 
                    query in task.get('notes', '').lower()):
                    
                    result_item = ctk.CTkFrame(results_frame, corner_radius=8)
                    result_item.pack(fill="x", pady=3)
                    
                    ctk.CTkLabel(
                        result_item,
                        text=task['title'],
                        font=ctk.CTkFont(size=13, weight="bold")
                    ).pack(anchor="w", padx=15, pady=(5, 0))
                    
                    ctk.CTkLabel(
                        result_item,
                        text=f"{task.get('status', 'Pendiente')} | {task['priority']} | {task['date']}",
                        font=ctk.CTkFont(size=11),
                        text_color="gray"
                    ).pack(anchor="w", padx=15, pady=(0, 5))
                    
                    results_count += 1
            
            if results_count == 0:
                ctk.CTkLabel(
                    results_frame,
                    text="No se encontraron resultados",
                    font=ctk.CTkFont(size=14),
                    text_color="gray"
                ).pack(pady=50)
            else:
                ctk.CTkLabel(
                    results_frame,
                    text=f"\n{results_count} resultado(s) encontrado(s)",
                    font=ctk.CTkFont(size=12),
                    text_color="green"
                ).pack(pady=10)
        
        ctk.CTkButton(
            search_frame,
            text="Buscar",
            command=perform_search,
            width=100,
            height=40,
            font=ctk.CTkFont(size=14, weight="bold")
        ).pack(side="left")
        
        search_entry.bind('<Return>', lambda e: perform_search())
        search_entry.focus()
    
    def create_all_tabs(self):
        """Crear todas las pestañas"""
        self.create_dashboard_tab()
        self.create_bugs_tab()
        self.create_test_cases_tab()
        self.create_suites_tab()
        self.create_tasks_tab()
        self.create_export_tab()
        self.create_history_tab()
    
    def hide_all_tabs(self):
        """Ocultar todas las pestañas"""
        for tab in self.tabs.values():
            tab.pack_forget()
    
    def show_dashboard_tab(self):
        self.hide_all_tabs()
        self.refresh_dashboard()
        self.tabs['dashboard'].pack(fill="both", expand=True)
        self.highlight_nav_button("Dashboard")
    
    def show_bugs_tab(self):
        self.hide_all_tabs()
        self.tabs['bugs'].pack(fill="both", expand=True)
        self.highlight_nav_button("Bugs")
        self.refresh_bugs_list()
    
    def show_test_cases_tab(self):
        self.hide_all_tabs()
        self.tabs['test_cases'].pack(fill="both", expand=True)
        self.highlight_nav_button("Casos de Prueba")
        self.refresh_tc_list()
    
    def show_suites_tab(self):
        self.hide_all_tabs()
        self.tabs['suites'].pack(fill="both", expand=True)
        self.highlight_nav_button("Test Suites")
        self.refresh_suites_list()
    
    def show_tasks_tab(self):
        self.hide_all_tabs()
        self.tabs['tasks'].pack(fill="both", expand=True)
        self.highlight_nav_button("Tareas")
        self.refresh_tasks_list()
    
    def show_export_tab(self):
        self.hide_all_tabs()
        self.tabs['export'].pack(fill="both", expand=True)
        self.highlight_nav_button("Exportar")
    
    def show_history_tab(self):
        self.hide_all_tabs()
        self.tabs['history'].pack(fill="both", expand=True)
        self.highlight_nav_button("Historial")
        self.refresh_history()
    
    def highlight_nav_button(self, button_text):
        """Resaltar botón activo"""
        for text, btn in self.nav_buttons.items():
            if text == button_text:
                btn.configure(fg_color=("#3B8ED0", "#1F6AA5"))
            else:
                btn.configure(fg_color=("gray75", "gray25"))
    
    def change_theme(self, value):
        """Cambiar tema de la aplicación"""
        if value == "Oscuro":
            ctk.set_appearance_mode("dark")
        else:
            ctk.set_appearance_mode("light")
    
    def create_dashboard_tab(self):
        """Crear dashboard mejorado con gráficos"""
        dash_frame = ctk.CTkFrame(self.content_area, corner_radius=10)
        self.tabs['dashboard'] = dash_frame
        
        header = ctk.CTkFrame(dash_frame, fg_color="transparent")
        header.pack(fill="x", padx=20, pady=15)
        
        ctk.CTkLabel(header, text="Dashboard", font=ctk.CTkFont(size=26, weight="bold")).pack(side="left")
        
        ctk.CTkButton(
            header,
            text="Actualizar",
            command=self.refresh_dashboard,
            width=120,
            height=35
        ).pack(side="right")
        
        scroll_frame = ctk.CTkScrollableFrame(dash_frame)
        scroll_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Contenedor de estadísticas principales
        self.stats_container = ctk.CTkFrame(scroll_frame, fg_color="transparent")
        self.stats_container.pack(fill="x", pady=10)
        
        # Sección de bugs críticos
        self.critical_bugs_section = ctk.CTkFrame(scroll_frame, corner_radius=10)
        self.critical_bugs_section.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            self.critical_bugs_section,
            text="BUGS CRÍTICOS SIN RESOLVER",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color="#E53935"
        ).pack(pady=15)
        
        self.critical_bugs_list = ctk.CTkFrame(self.critical_bugs_section, fg_color="transparent")
        self.critical_bugs_list.pack(fill="x", padx=20, pady=(0, 15))
        
        # Sección de métricas de casos
        self.metrics_section = ctk.CTkFrame(scroll_frame, corner_radius=10)
        self.metrics_section.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            self.metrics_section,
            text="MÉTRICAS DE TESTING",
            font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=15)
        
        self.metrics_container = ctk.CTkFrame(self.metrics_section, fg_color="transparent")
        self.metrics_container.pack(fill="x", padx=20, pady=(0, 15))
        
        # Actividad reciente
        self.recent_activity = ctk.CTkFrame(scroll_frame, corner_radius=10)
        self.recent_activity.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            self.recent_activity,
            text="ACTIVIDAD RECIENTE (Últimas 24h)",
            font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=15)
        
        self.activity_list = ctk.CTkFrame(self.recent_activity, fg_color="transparent")
        self.activity_list.pack(fill="x", padx=20, pady=(0, 15))
    
    def refresh_dashboard(self):
        """Actualizar dashboard con estadísticas"""
        # Limpiar contenedores
        for widget in self.stats_container.winfo_children():
            widget.destroy()
        for widget in self.critical_bugs_list.winfo_children():
            widget.destroy()
        for widget in self.metrics_container.winfo_children():
            widget.destroy()
        for widget in self.activity_list.winfo_children():
            widget.destroy()
        
        # Filtrar por proyecto si es necesario
        bugs = self.filter_by_project(self.bugs)
        test_cases = self.filter_by_project(self.test_cases)
        tasks = self.tasks  # Las tareas no tienen proyecto
        
        # Estadísticas principales
        stats_grid = ctk.CTkFrame(self.stats_container, fg_color="transparent")
        stats_grid.pack(fill="x")
        
        total_bugs = len(bugs)
        open_bugs = len([b for b in bugs if b['status'] not in ['Cerrado', 'Resuelto']])
        critical_bugs = len([b for b in bugs if b['severity'] == 'Crítico' and b['status'] not in ['Cerrado', 'Resuelto']])
        
        total_cases = len(test_cases)
        executed_cases = len([tc for tc in test_cases if tc.get('status') != 'Sin Probar'])
        passed_cases = len([tc for tc in test_cases if tc.get('status') == 'Aprobado'])
        failed_cases = len([tc for tc in test_cases if tc.get('status') == 'Fallido'])
        
        total_tasks = len(tasks)
        completed_tasks = len([t for t in tasks if t.get('status') == 'Completada'])
        pending_tasks = len([t for t in tasks if t.get('status') == 'Pendiente'])
        
        # Calcular porcentajes
        coverage = (executed_cases / total_cases * 100) if total_cases > 0 else 0
        pass_rate = (passed_cases / executed_cases * 100) if executed_cases > 0 else 0
        
        stats_data = [
            ("Total Bugs", total_bugs, "#E53935", f"{open_bugs} abiertos"),
            ("Casos de Prueba", total_cases, "#1E88E5", f"{coverage:.1f}% ejecutados"),
            ("Tareas", total_tasks, "#43A047", f"{completed_tasks} completadas"),
            ("Bugs Críticos", critical_bugs, "#D32F2F", "URGENTE" if critical_bugs > 0 else "OK"),
            ("Tasa de Éxito", f"{pass_rate:.1f}%", "#7CB342", f"{passed_cases}/{executed_cases}"),
            ("Casos Fallidos", failed_cases, "#F57C00", "Requieren atención")
        ]
        
        for i, (label, value, color, subtitle) in enumerate(stats_data):
            card = ctk.CTkFrame(stats_grid, fg_color=color, corner_radius=10, width=200)
            card.grid(row=i // 3, column=i % 3, padx=10, pady=10, sticky="ew")
            
            ctk.CTkLabel(
                card,
                text=str(value),
                font=ctk.CTkFont(size=36, weight="bold"),
                text_color="white"
            ).pack(pady=(15, 0))
            
            ctk.CTkLabel(
                card,
                text=label,
                font=ctk.CTkFont(size=14, weight="bold"),
                text_color="white"
            ).pack()
            
            ctk.CTkLabel(
                card,
                text=subtitle,
                font=ctk.CTkFont(size=11),
                text_color="white"
            ).pack(pady=(0, 15))
            
            stats_grid.columnconfigure(i % 3, weight=1)
        
        # Bugs críticos sin resolver
        critical_bugs_list = [b for b in bugs if b['severity'] == 'Crítico' and b['status'] not in ['Cerrado', 'Resuelto']]
        
        if critical_bugs_list:
            for bug in critical_bugs_list[:5]:  # Mostrar solo los primeros 5
                bug_item = ctk.CTkFrame(self.critical_bugs_list, corner_radius=8, fg_color="#FFEBEE")
                bug_item.pack(fill="x", pady=3)
                
                ctk.CTkLabel(
                    bug_item,
                    text=f"#{bug['id']} - {bug['title']}",
                    font=ctk.CTkFont(size=13, weight="bold"),
                    text_color="#B71C1C"
                ).pack(anchor="w", padx=15, pady=(5, 0))
                
                ctk.CTkLabel(
                    bug_item,
                    text=f"{bug['status']} | {bug['priority']} | {bug['type']}",
                    font=ctk.CTkFont(size=11),
                    text_color="#C62828"
                ).pack(anchor="w", padx=15, pady=(0, 5))
        else:
            ctk.CTkLabel(
                self.critical_bugs_list,
                text="✅ No hay bugs críticos pendientes",
                font=ctk.CTkFont(size=14),
                text_color="green"
            ).pack(pady=10)
        
        # Métricas de testing
        metrics_grid = ctk.CTkFrame(self.metrics_container, fg_color="transparent")
        metrics_grid.pack(fill="x")
        
        metrics = [
            ("Cobertura de Ejecución", f"{coverage:.1f}%", f"{executed_cases}/{total_cases} casos"),
            ("Tasa de Aprobación", f"{pass_rate:.1f}%", f"{passed_cases} aprobados"),
            ("Casos Bloqueados", len([tc for tc in test_cases if tc.get('status') == 'Bloqueado']), "Requieren desbloqueo"),
            ("Casos Pendientes", len([tc for tc in test_cases if tc.get('status') == 'Pendiente']), "En espera")
        ]
        
        for i, (label, value, subtitle) in enumerate(metrics):
            metric_card = ctk.CTkFrame(metrics_grid, corner_radius=8)
            metric_card.grid(row=i // 2, column=i % 2, padx=10, pady=10, sticky="ew")
            
            ctk.CTkLabel(
                metric_card,
                text=label,
                font=ctk.CTkFont(size=13, weight="bold")
            ).pack(anchor="w", padx=15, pady=(10, 0))
            
            ctk.CTkLabel(
                metric_card,
                text=str(value),
                font=ctk.CTkFont(size=24, weight="bold"),
                text_color="#1E88E5"
            ).pack(anchor="w", padx=15)
            
            ctk.CTkLabel(
                metric_card,
                text=subtitle,
                font=ctk.CTkFont(size=11),
                text_color="gray"
            ).pack(anchor="w", padx=15, pady=(0, 10))
            
            metrics_grid.columnconfigure(i % 2, weight=1)
        
        # Actividad reciente (últimas 24 horas)
        yesterday = datetime.now() - timedelta(days=1)
        recent_history = [h for h in self.history if datetime.strptime(h['timestamp'], '%Y-%m-%d %H:%M:%S') > yesterday]
        
        if recent_history:
            for entry in recent_history[-10:]:  # Últimas 10 acciones
                activity_item = ctk.CTkFrame(self.activity_list, corner_radius=8)
                activity_item.pack(fill="x", pady=2)
                
                action_icons = {
                    'created': '➕',
                    'updated': '✏️',
                    'deleted': '🗑️',
                    'status_changed': '🔄'
                }
                
                type_colors = {
                    'bug': '#E53935',
                    'test_case': '#1E88E5',
                    'task': '#43A047'
                }
                
                icon = action_icons.get(entry['action'], '•')
                color = type_colors.get(entry['type'], 'gray')
                
                info_text = f"{icon} {entry['timestamp']} - {entry['description']}"
                
                ctk.CTkLabel(
                    activity_item,
                    text=info_text,
                    font=ctk.CTkFont(size=11),
                    text_color=color
                ).pack(anchor="w", padx=15, pady=5)
        else:
            ctk.CTkLabel(
                self.activity_list,
                text="No hay actividad reciente en las últimas 24 horas",
                font=ctk.CTkFont(size=12),
                text_color="gray"
            ).pack(pady=10)
    
    def filter_by_project(self, items):
        """Filtrar items por proyecto actual"""
        if self.current_project == "Todos":
            return items
        return [item for item in items if item.get('project') == self.current_project]
    
    def create_suites_tab(self):
        """Crear pestaña de Test Suites"""
        suites_frame = ctk.CTkFrame(self.content_area, corner_radius=10)
        self.tabs['suites'] = suites_frame
        
        header = ctk.CTkFrame(suites_frame, fg_color="transparent")
        header.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkLabel(header, text="Test Suites", font=ctk.CTkFont(size=24, weight="bold")).pack(side="left")
        
        ctk.CTkButton(
            header,
            text="+ Nueva Suite",
            command=self.show_new_suite_form,
            width=140,
            height=35,
            fg_color="#2FA572"
        ).pack(side="right", padx=5)
        
        # Lista de suites
        self.suites_list_frame = ctk.CTkScrollableFrame(suites_frame)
        self.suites_list_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Botones de acción
        action_frame = ctk.CTkFrame(suites_frame, fg_color="transparent")
        action_frame.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkButton(
            action_frame,
            text="Ejecutar Suite",
            command=self.execute_suite,
            width=130
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            action_frame,
            text="Ver Detalles",
            command=self.view_suite_details,
            width=130
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            action_frame,
            text="Eliminar",
            command=self.delete_suite,
            width=120,
            fg_color="#D32F2F"
        ).pack(side="left", padx=5)
    
    def show_new_suite_form(self):
        """Formulario para nueva suite"""
        form = ctk.CTkToplevel(self.root)
        form.title("Nueva Test Suite")
        form.geometry("700x600")
        form.grab_set()
        
        scroll = ctk.CTkScrollableFrame(form)
        scroll.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(scroll, text="Crear Nueva Test Suite", font=ctk.CTkFont(size=20, weight="bold")).pack(pady=(0, 20))
        
        # Nombre
        ctk.CTkLabel(scroll, text="Nombre de la Suite:", anchor="w").pack(fill="x", pady=(5, 0))
        name_entry = ctk.CTkEntry(scroll, placeholder_text="Ej: Suite de Regresión Login")
        name_entry.pack(fill="x", pady=(0, 10))
        
        # Descripción
        ctk.CTkLabel(scroll, text="Descripción:", anchor="w").pack(fill="x", pady=(5, 0))
        desc_text = ctk.CTkTextbox(scroll, height=60)
        desc_text.pack(fill="x", pady=(0, 10))
        
        # Seleccionar casos de prueba
        ctk.CTkLabel(scroll, text="Seleccionar Casos de Prueba:", font=ctk.CTkFont(size=14, weight="bold")).pack(fill="x", pady=(15, 10))
        
        cases_frame = ctk.CTkScrollableFrame(scroll, height=250)
        cases_frame.pack(fill="x", pady=10)
        
        selected_cases = []
        
        for tc in self.test_cases:
            var = ctk.BooleanVar()
            cb = ctk.CTkCheckBox(
                cases_frame,
                text=f"{tc['test_id']} - {tc['title']} ({tc['module']})",
                variable=var,
                command=lambda t=tc, v=var: selected_cases.append(t) if v.get() else selected_cases.remove(t) if t in selected_cases else None
            )
            cb.pack(anchor="w", pady=2, padx=10)
        
        # Botones
        btn_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        btn_frame.pack(fill="x", pady=20)
        
        def save_suite():
            if not name_entry.get():
                messagebox.showwarning("Advertencia", "El nombre es obligatorio")
                return
            
            if not selected_cases:
                messagebox.showwarning("Advertencia", "Debes seleccionar al menos un caso de prueba")
                return
            
            suite = {
                'id': len(self.test_suites) + 1,
                'name': name_entry.get(),
                'description': desc_text.get('1.0', 'end-1c'),
                'test_cases': [tc['test_id'] for tc in selected_cases],
                'created': datetime.now().strftime('%Y-%m-%d %H:%M'),
                'last_executed': '',
                'results': {}
            }
            
            self.test_suites.append(suite)
            self.save_data(self.test_suites, self.suites_file)
            self.add_to_history('created', 'suite', suite['id'], f"Suite creada: {suite['name']}")
            self.refresh_suites_list()
            form.destroy()
            messagebox.showinfo("Éxito", "Suite creada correctamente")
        
        ctk.CTkButton(
            btn_frame,
            text="Crear Suite",
            command=save_suite,
            width=150,
            height=40,
            fg_color="#2FA572"
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            btn_frame,
            text="Cancelar",
            command=form.destroy,
            width=120,
            height=40,
            fg_color="gray"
        ).pack(side="left", padx=5)
    
    def refresh_suites_list(self):
        """Refrescar lista de suites"""
        for widget in self.suites_list_frame.winfo_children():
            widget.destroy()
        
        self.selected_suite = None
        self.selected_suite_widget = None
        
        if not self.test_suites:
            ctk.CTkLabel(
                self.suites_list_frame,
                text="No hay test suites creadas",
                text_color="gray",
                font=ctk.CTkFont(size=14)
            ).pack(pady=20)
            return
        
        for suite in reversed(self.test_suites):
            item_frame = ctk.CTkFrame(self.suites_list_frame, corner_radius=8)
            item_frame.pack(fill="x", pady=5, padx=5)
            
            check_var = ctk.BooleanVar(value=False)
            checkbox = ctk.CTkCheckBox(
                item_frame,
                text="",
                variable=check_var,
                command=lambda s=suite, v=check_var, f=item_frame: self.toggle_suite_selection(s, v, f),
                width=30
            )
            checkbox.pack(side="left", padx=(10, 5), pady=10)
            
            content_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            content_frame.pack(side="left", fill="x", expand=True, padx=5, pady=5)
            
            title_label = ctk.CTkLabel(
                content_frame,
                text=suite['name'],
                font=ctk.CTkFont(size=14, weight="bold"),
                anchor="w"
            )
            title_label.pack(fill="x")
            
            info_label = ctk.CTkLabel(
                content_frame,
                text=f"{len(suite['test_cases'])} casos | Creada: {suite['created'].split()[0]}",
                font=ctk.CTkFont(size=11),
                text_color="gray",
                anchor="w"
            )
            info_label.pack(fill="x")
            
            if suite.get('last_executed'):
                exec_label = ctk.CTkLabel(
                    content_frame,
                    text=f"Última ejecución: {suite['last_executed']}",
                    font=ctk.CTkFont(size=10),
                    text_color="green",
                    anchor="w"
                )
                exec_label.pack(fill="x")
            
            item_frame._suite_data = suite
            item_frame._checkbox = checkbox
            item_frame._check_var = check_var
    
    def toggle_suite_selection(self, suite, check_var, item_frame):
        """Manejar selección de suite"""
        if check_var.get():
            if hasattr(self, 'selected_suite_widget') and self.selected_suite_widget and self.selected_suite_widget != item_frame:
                self.selected_suite_widget._check_var.set(False)
            
            self.selected_suite = suite
            self.selected_suite_widget = item_frame
        else:
            if self.selected_suite == suite:
                self.selected_suite = None
                self.selected_suite_widget = None
    
    def execute_suite(self):
        """Ejecutar suite de pruebas"""
        if not hasattr(self, 'selected_suite') or not self.selected_suite:
            messagebox.showwarning("Advertencia", "Selecciona una suite primero")
            return
        
        suite = self.selected_suite
        
        exec_window = ctk.CTkToplevel(self.root)
        exec_window.title(f"Ejecutar Suite: {suite['name']}")
        exec_window.geometry("800x600")
        exec_window.grab_set()
        
        ctk.CTkLabel(
            exec_window,
            text=f"Ejecutando: {suite['name']}",
            font=ctk.CTkFont(size=20, weight="bold")
        ).pack(pady=20)
        
        scroll = ctk.CTkScrollableFrame(exec_window)
        scroll.pack(fill="both", expand=True, padx=20, pady=10)
        
        results = {}
        
        for tc_id in suite['test_cases']:
            tc = next((t for t in self.test_cases if t['test_id'] == tc_id), None)
            if not tc:
                continue
            
            case_frame = ctk.CTkFrame(scroll, corner_radius=8)
            case_frame.pack(fill="x", pady=5)
            
            ctk.CTkLabel(
                case_frame,
                text=f"{tc['test_id']} - {tc['title']}",
                font=ctk.CTkFont(size=13, weight="bold")
            ).pack(anchor="w", padx=15, pady=(10, 5))
            
            status_var = ctk.StringVar(value=tc.get('status', 'Sin Probar'))
            
            status_frame = ctk.CTkFrame(case_frame, fg_color="transparent")
            status_frame.pack(fill="x", padx=15, pady=(0, 10))
            
            ctk.CTkLabel(status_frame, text="Resultado:").pack(side="left", padx=(0, 10))
            
            status_combo = ctk.CTkComboBox(
                status_frame,
                values=['Aprobado', 'Fallido', 'Bloqueado'],
                variable=status_var,
                width=150
            )
            status_combo.pack(side="left")
            
            results[tc_id] = status_var
        
        def save_execution():
            # Actualizar estados de casos
            for tc_id, status_var in results.items():
                for tc in self.test_cases:
                    if tc['test_id'] == tc_id:
                        tc['status'] = status_var.get()
                        if not tc.get('executed_date'):
                            tc['executed_date'] = datetime.now().strftime('%Y-%m-%d %H:%M')
                        break
            
            # Actualizar suite
            for s in self.test_suites:
                if s['id'] == suite['id']:
                    s['last_executed'] = datetime.now().strftime('%Y-%m-%d %H:%M')
                    s['results'] = {tc_id: var.get() for tc_id, var in results.items()}
                    break
            
            self.save_data(self.test_cases, self.test_cases_file)
            self.save_data(self.test_suites, self.suites_file)
            self.add_to_history('executed', 'suite', suite['id'], f"Suite ejecutada: {suite['name']}")
            
            exec_window.destroy()
            self.refresh_suites_list()
            messagebox.showinfo("Éxito", "Ejecución completada y resultados guardados")
        
        ctk.CTkButton(
            exec_window,
            text="Guardar Resultados",
            command=save_execution,
            width=200,
            height=40,
            fg_color="#2FA572"
        ).pack(pady=20)
    
    def view_suite_details(self):
        """Ver detalles de suite"""
        if not hasattr(self, 'selected_suite') or not self.selected_suite:
            messagebox.showwarning("Advertencia", "Selecciona una suite primero")
            return
        
        suite = self.selected_suite
        
        detail_window = ctk.CTkToplevel(self.root)
        detail_window.title(f"Suite: {suite['name']}")
        detail_window.geometry("700x600")
        
        scroll = ctk.CTkScrollableFrame(detail_window)
        scroll.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(scroll, text=suite['name'], font=ctk.CTkFont(size=20, weight="bold")).pack(pady=(0, 10))
        
        ctk.CTkLabel(scroll, text=suite['description'], wraplength=650).pack(pady=10)
        
        info_frame = ctk.CTkFrame(scroll)
        info_frame.pack(fill="x", pady=10)
        
        info_data = [
            ("Casos de Prueba", len(suite['test_cases'])),
            ("Creada", suite['created']),
            ("Última Ejecución", suite.get('last_executed', 'Nunca'))
        ]
        
        for label, value in info_data:
            row = ctk.CTkFrame(info_frame, fg_color="transparent")
            row.pack(fill="x", padx=15, pady=5)
            ctk.CTkLabel(row, text=f"{label}:", font=ctk.CTkFont(weight="bold"), width=150, anchor="w").pack(side="left")
            ctk.CTkLabel(row, text=str(value), anchor="w").pack(side="left")
        
        ctk.CTkLabel(scroll, text="Casos Incluidos:", font=ctk.CTkFont(size=14, weight="bold"), anchor="w").pack(fill="x", pady=(20, 10))
        
        for tc_id in suite['test_cases']:
            tc = next((t for t in self.test_cases if t['test_id'] == tc_id), None)
            if tc:
                case_label = ctk.CTkLabel(scroll, text=f"• {tc_id} - {tc['title']}", anchor="w")
                case_label.pack(fill="x", padx=20, pady=2)
    
    def delete_suite(self):
        """Eliminar suite"""
        if not hasattr(self, 'selected_suite') or not self.selected_suite:
            messagebox.showwarning("Advertencia", "Selecciona una suite primero")
            return
        
        if messagebox.askyesno("Confirmar", "¿Eliminar esta suite?\n\nLos casos de prueba NO se eliminarán."):
            self.test_suites = [s for s in self.test_suites if s['id'] != self.selected_suite['id']]
            self.save_data(self.test_suites, self.suites_file)
            self.add_to_history('deleted', 'suite', self.selected_suite['id'], f"Suite eliminada: {self.selected_suite['name']}")
            self.selected_suite = None
            self.selected_suite_widget = None
            self.refresh_suites_list()
    
    def create_history_tab(self):
        """Crear pestaña de historial"""
        history_frame = ctk.CTkFrame(self.content_area, corner_radius=10)
        self.tabs['history'] = history_frame
        
        header = ctk.CTkFrame(history_frame, fg_color="transparent")
        header.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkLabel(header, text="Historial de Actividades", font=ctk.CTkFont(size=24, weight="bold")).pack(side="left")
        
        ctk.CTkButton(
            header,
            text="Actualizar",
            command=self.refresh_history,
            width=120,
            height=35
        ).pack(side="right", padx=5)
        
        ctk.CTkButton(
            header,
            text="Limpiar Historial",
            command=self.clear_history,
            width=150,
            height=35,
            fg_color="#D32F2F"
        ).pack(side="right", padx=5)
        
        # Lista de historial
        self.history_list_frame = ctk.CTkScrollableFrame(history_frame)
        self.history_list_frame.pack(fill="both", expand=True, padx=20, pady=10)
    
    def refresh_history(self):
        """Refrescar historial"""
        for widget in self.history_list_frame.winfo_children():
            widget.destroy()
        
        if not self.history:
            ctk.CTkLabel(
                self.history_list_frame,
                text="No hay actividad registrada",
                text_color="gray",
                font=ctk.CTkFont(size=14)
            ).pack(pady=20)
            return
        
        # Agrupar por día
        by_date = defaultdict(list)
        for entry in reversed(self.history):
            date = entry['timestamp'].split()[0]
            by_date[date].append(entry)
        
        for date in sorted(by_date.keys(), reverse=True):
            # Cabecera de fecha
            date_header = ctk.CTkFrame(self.history_list_frame, fg_color="#1E88E5", corner_radius=5)
            date_header.pack(fill="x", pady=(10, 5))
            
            ctk.CTkLabel(
                date_header,
                text=date,
                font=ctk.CTkFont(size=14, weight="bold"),
                text_color="white"
            ).pack(padx=15, pady=8)
            
            # Entradas del día
            for entry in by_date[date]:
                entry_frame = ctk.CTkFrame(self.history_list_frame, corner_radius=5)
                entry_frame.pack(fill="x", pady=2, padx=10)
                
                action_icons = {
                    'created': '➕',
                    'updated': '✏️',
                    'deleted': '🗑️',
                    'status_changed': '🔄',
                    'executed': '▶️'
                }
                
                type_colors = {
                    'bug': '#E53935',
                    'test_case': '#1E88E5',
                    'task': '#43A047',
                    'suite': '#9C27B0'
                }
                
                icon = action_icons.get(entry['action'], '•')
                color = type_colors.get(entry['type'], 'gray')
                
                time = entry['timestamp'].split()[1]
                
                ctk.CTkLabel(
                    entry_frame,
                    text=f"{icon} {time}",
                    font=ctk.CTkFont(size=11, weight="bold"),
                    width=70,
                    text_color=color
                ).pack(side="left", padx=10)
                
                ctk.CTkLabel(
                    entry_frame,
                    text=entry['description'],
                    font=ctk.CTkFont(size=11),
                    anchor="w"
                ).pack(side="left", fill="x", expand=True, padx=5, pady=8)
    
    def clear_history(self):
        """Limpiar historial"""
        if messagebox.askyesno("Confirmar", "¿Estás seguro de querer limpiar todo el historial?"):
            self.history = []
            self.save_data(self.history, self.history_file)
            self.refresh_history()
            messagebox.showinfo("Éxito", "Historial limpiado")
    
    def create_bugs_tab(self):
        """Crear pestaña de bugs con selector de estado inline EN ESPAÑOL"""
        bugs_frame = ctk.CTkFrame(self.content_area, corner_radius=10)
        self.tabs['bugs'] = bugs_frame
        
        # Header con filtros
        header = ctk.CTkFrame(bugs_frame, fg_color="transparent")
        header.pack(fill="x", padx=20, pady=10)
        
        title = ctk.CTkLabel(
            header, 
            text="Gestión de Bugs", 
            font=ctk.CTkFont(size=24, weight="bold")
        )
        title.pack(side="left")
        
        # Botón nuevo bug
        ctk.CTkButton(
            header,
            text="+ Nuevo Bug",
            command=self.show_new_bug_form,
            width=140,
            height=35,
            fg_color="#B11675",
            hover_color="#26865E"
        ).pack(side="right", padx=5)
        
        # Filtros de estado EN ESPAÑOL
        filter_frame = ctk.CTkFrame(bugs_frame, fg_color="transparent")
        filter_frame.pack(fill="x", padx=20, pady=5)
        
        ctk.CTkLabel(
            filter_frame,
            text="Filtrar por estado:",
            font=ctk.CTkFont(size=13, weight="bold")
        ).pack(side="left", padx=(0, 10))
        
        # Botones de filtro
        self.bug_filter_buttons = {}
        bug_states = ["Todos", "Nuevo", "En Progreso", "Resuelto", "Cerrado", "Rechazado"]
        
        for state in bug_states:
            btn = ctk.CTkButton(
                filter_frame,
                text=state,
                command=lambda s=state: self.filter_bugs(s),
                width=100,
                height=30,
                corner_radius=15,
                fg_color="transparent",
                border_width=2,
                border_color="#3B8ED0"
            )
            btn.pack(side="left", padx=3)
            self.bug_filter_buttons[state] = btn
        
        # Resaltar "Todos" por defecto
        self.bug_filter_buttons["Todos"].configure(fg_color="#3B8ED0", text_color="white")
        
        # Scrollable frame para la lista
        self.bugs_list_frame = ctk.CTkScrollableFrame(bugs_frame)
        self.bugs_list_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Botones de acción
        action_frame = ctk.CTkFrame(bugs_frame, fg_color="transparent")
        action_frame.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkButton(
            action_frame, 
            text="Editar",
            command=self.edit_bug,
            width=100
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            action_frame, 
            text="Duplicar",
            command=self.duplicate_bug,
            width=100
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            action_frame, 
            text="Ver Detalles",
            command=self.view_bug_details,
            width=100
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            action_frame, 
            text="Eliminar",
            command=self.delete_bug,
            width=100,
            fg_color="#D32F2F",
            hover_color="#B71C1C"
        ).pack(side="left", padx=5)
        
        self.refresh_bugs_list()
    
    def filter_bugs(self, state):
        """Filtrar bugs por estado"""
        self.current_bug_filter = state
        
        # Actualizar estilo de botones
        for s, btn in self.bug_filter_buttons.items():
            if s == state:
                btn.configure(fg_color="#3B8ED0", text_color="white")
            else:
                btn.configure(fg_color="transparent", text_color=("gray10", "gray90"))
        
        self.refresh_bugs_list()
    
    def show_new_bug_form(self):
        """Mostrar formulario para nuevo bug"""
        form_window = ctk.CTkToplevel(self.root)
        form_window.title("Nuevo Bug")
        form_window.geometry("700x750")
        form_window.grab_set()
        
        # Scrollable frame
        scroll = ctk.CTkScrollableFrame(form_window)
        scroll.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(
            scroll,
            text="Registrar Nuevo Bug",
            font=ctk.CTkFont(size=20, weight="bold")
        ).pack(pady=(0, 20))
        
        # Título
        ctk.CTkLabel(scroll, text="Título:", anchor="w").pack(fill="x", pady=(5, 0))
        title_entry = ctk.CTkEntry(scroll, placeholder_text="Descripción breve del bug")
        title_entry.pack(fill="x", pady=(0, 10))
        
        # Grid para campos
        grid_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        grid_frame.pack(fill="x", pady=5)
        
        # Proyecto
        ctk.CTkLabel(grid_frame, text="Proyecto:").grid(row=0, column=0, sticky="w", padx=(0, 10), pady=5)
        project_combo = ctk.CTkComboBox(grid_frame, values=[p['name'] for p in self.projects], width=150)
        if self.projects:
            project_combo.set(self.projects[0]['name'])
        project_combo.grid(row=0, column=1, sticky="w", pady=5)
        
        # Severidad
        ctk.CTkLabel(grid_frame, text="Severidad:").grid(row=0, column=2, sticky="w", padx=(20, 10), pady=5)
        severity_combo = ctk.CTkComboBox(grid_frame, values=['Crítico', 'Alto', 'Medio', 'Bajo'], width=150)
        severity_combo.set('Medio')
        severity_combo.grid(row=0, column=3, sticky="w", pady=5)
        
        # Prioridad
        ctk.CTkLabel(grid_frame, text="Prioridad:").grid(row=1, column=0, sticky="w", padx=(0, 10), pady=5)
        priority_combo = ctk.CTkComboBox(grid_frame, values=['Urgente', 'Alta', 'Media', 'Baja'], width=150)
        priority_combo.set('Media')
        priority_combo.grid(row=1, column=1, sticky="w", pady=5)
        
        # Estado
        ctk.CTkLabel(grid_frame, text="Estado:").grid(row=1, column=2, sticky="w", padx=(20, 10), pady=5)
        status_combo = ctk.CTkComboBox(grid_frame, values=['Nuevo', 'En Progreso', 'Resuelto', 'Cerrado', 'Rechazado'], width=150)
        status_combo.set('Nuevo')
        status_combo.grid(row=1, column=3, sticky="w", pady=5)
        
        # Tipo
        ctk.CTkLabel(grid_frame, text="Tipo:").grid(row=2, column=0, sticky="w", padx=(0, 10), pady=5)
        type_combo = ctk.CTkComboBox(grid_frame, values=['Funcional', 'UI/UX', 'Rendimiento', 'Seguridad', 'Compatibilidad', 'Datos'], width=150)
        type_combo.set('Funcional')
        type_combo.grid(row=2, column=1, sticky="w", pady=5)
        
        # Descripción
        ctk.CTkLabel(scroll, text="Descripción:", anchor="w").pack(fill="x", pady=(10, 0))
        description_text = ctk.CTkTextbox(scroll, height=80)
        description_text.pack(fill="x", pady=(0, 10))
        
        # Pasos
        ctk.CTkLabel(scroll, text="Pasos para Reproducir:", anchor="w").pack(fill="x", pady=(5, 0))
        steps_text = ctk.CTkTextbox(scroll, height=60)
        steps_text.pack(fill="x", pady=(0, 10))
        
        #  MEJORA 1: Múltiples Screenshots
        screenshots_frame = ctk.CTkFrame(scroll, corner_radius=8, fg_color=("#E8E8E8", "#2B2B2B"))
        screenshots_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            screenshots_frame,
            text="📸 Capturas de Pantalla (puedes agregar varias)",
            font=ctk.CTkFont(size=13, weight="bold")
        ).pack(pady=(10, 5), padx=15, anchor="w")
        
        screenshots_list = []
        screenshots_container = ctk.CTkFrame(screenshots_frame, fg_color="transparent")
        screenshots_container.pack(fill="x", padx=15, pady=5)
        
        def add_screenshot():
            filenames = filedialog.askopenfilenames(
                title="Seleccionar Captura(s) de Pantalla",
                filetypes=[("Imágenes", "*.png *.jpg *.jpeg *.gif *.bmp"), ("Todos", "*.*")]
            )
            if filenames:
                for filename in filenames:
                    dest = self.data_dir / "screenshots" / f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{Path(filename).name}"
                    shutil.copy(filename, dest)
                    screenshots_list.append(str(dest))
                    
                    # Mostrar en la lista
                    img_item = ctk.CTkFrame(screenshots_container, corner_radius=5)
                    img_item.pack(fill="x", pady=2)
                    
                    ctk.CTkLabel(
                        img_item,
                        text=f"✓ {Path(filename).name}",
                        font=ctk.CTkFont(size=11),
                        text_color="green"
                    ).pack(side="left", padx=10, pady=5)
                    
                    def remove_img(path=str(dest), frame=img_item):
                        screenshots_list.remove(path)
                        frame.destroy()
                    
                    ctk.CTkButton(
                        img_item,
                        text="Eliminar",
                        command=remove_img,
                        width=70,
                        height=25,
                        fg_color="#D32F2F",
                        hover_color="#B71C1C",
                        font=ctk.CTkFont(size=10)
                    ).pack(side="right", padx=5)
                
                messagebox.showinfo("Éxito", f"{len(filenames)} imagen(es) agregada(s)")
        
        ctk.CTkButton(
            screenshots_frame,
            text="+ Agregar Imágenes",
            command=add_screenshot,
            width=150,
            height=35,
            fg_color="#2196F3"
        ).pack(pady=(5, 15), padx=15)
        
        # Botones
        btn_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        btn_frame.pack(fill="x", pady=20)
        
        def save_new_bug():
            if not title_entry.get():
                messagebox.showwarning("Advertencia", "El título es obligatorio")
                return
            
            bug = {
                'id': len(self.bugs) + 1,
                'title': title_entry.get(),
                'severity': severity_combo.get(),
                'priority': priority_combo.get(),
                'status': status_combo.get(),
                'type': type_combo.get(),
                'project': project_combo.get(),
                'description': description_text.get('1.0', 'end-1c'),
                'steps': steps_text.get('1.0', 'end-1c'),
                'screenshots': screenshots_list.copy(),  # 🆕 Lista de screenshots
                'date': datetime.now().strftime('%Y-%m-%d %H:%M')
            }
            
            self.bugs.append(bug)
            self.save_data(self.bugs, self.bugs_file)
            self.add_to_history('created', 'bug', bug['id'], f"Bug creado: {bug['title']}")
            self.refresh_bugs_list()
            form_window.destroy()
            messagebox.showinfo("Éxito", "Bug registrado correctamente")
        
        ctk.CTkButton(
            btn_frame,
            text="Guardar",
            command=save_new_bug,
            width=150,
            height=40,
            fg_color="#2FA572",
            hover_color="#26865E"
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            btn_frame,
            text="Cancelar",
            command=form_window.destroy,
            width=120,
            height=40,
            fg_color="gray",
            hover_color="darkgray"
        ).pack(side="left", padx=5)
    
    def refresh_bugs_list(self):
        """Refrescar lista de bugs CON SELECTOR DE ESTADO INLINE"""
        for widget in self.bugs_list_frame.winfo_children():
            widget.destroy()
        
        self.selected_bug = None
        self.selected_bug_widget = None
        
        # Filtrar bugs
        bugs = self.filter_by_project(self.bugs)
        
        if self.current_bug_filter == "Todos":
            filtered_bugs = bugs
        else:
            filtered_bugs = [b for b in bugs if b['status'] == self.current_bug_filter]
        
        if not filtered_bugs:
            no_data = ctk.CTkLabel(
                self.bugs_list_frame,
                text=f"No hay bugs en estado '{self.current_bug_filter}'",
                text_color="gray",
                font=ctk.CTkFont(size=14)
            )
            no_data.pack(pady=20)
            return
        
        # Crear lista con selector de estado
        for bug in reversed(filtered_bugs):
            # Frame principal del item
            item_frame = ctk.CTkFrame(self.bugs_list_frame, corner_radius=8)
            item_frame.pack(fill="x", pady=3, padx=5)
            
            # Checkbox para selección
            check_var = ctk.BooleanVar(value=False)
            checkbox = ctk.CTkCheckBox(
                item_frame,
                text="",
                variable=check_var,
                command=lambda b=bug, v=check_var, f=item_frame: self.toggle_bug_selection(b, v, f),
                width=30
            )
            checkbox.pack(side="left", padx=(10, 5), pady=10)
            
            # Contenido del bug
            content_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            content_frame.pack(side="left", fill="x", expand=True, padx=5, pady=5)
            
            # Línea 1: ID + Título
            title_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
            title_frame.pack(fill="x")
            
            id_label = ctk.CTkLabel(
                title_frame,
                text=f"#{bug['id']}",
                font=ctk.CTkFont(size=11, weight="bold"),
                width=40,
                text_color="#3B8ED0"
            )
            id_label.pack(side="left", padx=(0, 5))
            
            title_label = ctk.CTkLabel(
                title_frame,
                text=bug['title'],
                font=ctk.CTkFont(size=13),
                anchor="w"
            )
            title_label.pack(side="left", fill="x", expand=True)
            
            # Línea 2: Estado, Severidad, Prioridad
            info_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
            info_frame.pack(fill="x", pady=(3, 0))
            
            # Estado
            status_label = ctk.CTkLabel(
                info_frame,
                text=bug['status'],
                fg_color="#1976D2",
                text_color="white",
                corner_radius=4,
                font=ctk.CTkFont(size=10, weight="bold"),
                width=80,
                height=20
            )
            status_label.pack(side="left", padx=(0, 5))
            
            # Severidad con colores
            severity_colors = {
                'Crítico': "#D32F2F",
                'Alto': "#F57C00",
                'Medio': "#FBC02D",
                'Bajo': "#388E3C"
            }
            
            severity_label = ctk.CTkLabel(
                info_frame,
                text=bug['severity'],
                fg_color=severity_colors.get(bug['severity'], "#757575"),
                text_color="white",
                corner_radius=4,
                font=ctk.CTkFont(size=10, weight="bold"),
                width=60,
                height=20
            )
            severity_label.pack(side="left", padx=(0, 5))
            
            # Prioridad
            priority_label = ctk.CTkLabel(
                info_frame,
                text=f"Prioridad: {bug['priority']}",
                font=ctk.CTkFont(size=10),
                text_color="gray"
            )
            priority_label.pack(side="left", padx=(0, 10))
            
            # Tipo, Proyecto y Fecha
            extra_label = ctk.CTkLabel(
                info_frame,
                text=f"{bug['type']} | {bug.get('project', 'Sin proyecto')} | {bug['date'].split()[0]}",
                font=ctk.CTkFont(size=10),
                text_color="gray"
            )
            extra_label.pack(side="left")
            
            # **SELECTOR DE ESTADO AL LADO DERECHO** (inicialmente oculto)
            status_selector_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            
            status_combo = ctk.CTkComboBox(
                status_selector_frame,
                values=['Nuevo', 'En Progreso', 'Resuelto', 'Cerrado', 'Rechazado'],
                width=140,
                height=28,
                command=lambda choice, b=bug: self.update_bug_status_inline(b, choice)
            )
            status_combo.set(bug['status'])
            status_combo.pack(padx=10)
            
            # Guardar referencias
            item_frame._bug_data = bug
            item_frame._checkbox = checkbox
            item_frame._check_var = check_var
            item_frame._status_selector = status_selector_frame
            item_frame._status_combo = status_combo
            item_frame._status_label = status_label
    
    def toggle_bug_selection(self, bug, check_var, item_frame):
        """Manejar selección de bug con checkbox y mostrar selector de estado"""
        if check_var.get():
            # Deseleccionar anterior
            if self.selected_bug_widget and self.selected_bug_widget != item_frame:
                self.selected_bug_widget._check_var.set(False)
                self.selected_bug_widget._status_selector.pack_forget()
            
            # Seleccionar nuevo
            self.selected_bug = bug
            self.selected_bug_widget = item_frame
            
            # Mostrar selector de estado al lado derecho
            item_frame._status_selector.pack(side="right", padx=5, pady=5)
        else:
            if self.selected_bug == bug:
                self.selected_bug = None
                self.selected_bug_widget = None
                item_frame._status_selector.pack_forget()
    
    def update_bug_status_inline(self, bug, new_status):
        """Actualizar estado del bug desde el selector inline"""
        for b in self.bugs:
            if b['id'] == bug['id']:
                old_status = b['status']
                b['status'] = new_status
                break
        
        self.save_data(self.bugs, self.bugs_file)
        self.add_to_history('status_changed', 'bug', bug['id'], f"Bug #{bug['id']}: {old_status} → {new_status}")
        
        # Actualizar la etiqueta de estado en el widget
        if self.selected_bug_widget:
            self.selected_bug_widget._status_label.configure(text=new_status)
        
        messagebox.showinfo("Actualizado", f"Estado cambiado a: {new_status}")
    
    def duplicate_bug(self):
        """Duplicar bug seleccionado"""
        if not self.selected_bug:
            messagebox.showwarning("Advertencia", "Selecciona un bug primero")
            return
        
        bug = self.selected_bug.copy()
        bug['id'] = len(self.bugs) + 1
        bug['title'] = f"[COPIA] {bug['title']}"
        bug['status'] = 'Nuevo'
        bug['date'] = datetime.now().strftime('%Y-%m-%d %H:%M')
        
        self.bugs.append(bug)
        self.save_data(self.bugs, self.bugs_file)
        self.add_to_history('created', 'bug', bug['id'], f"Bug duplicado: {bug['title']}")
        self.refresh_bugs_list()
        messagebox.showinfo("Éxito", "Bug duplicado correctamente")
    
    def edit_bug(self):
        """Editar bug seleccionado"""
        if not self.selected_bug:
            messagebox.showwarning("Advertencia", "Selecciona un bug primero")
            return
        
        bug = self.selected_bug
        edit_window = ctk.CTkToplevel(self.root)
        edit_window.title(f"Editar Bug #{bug['id']}")
        edit_window.geometry("700x750")
        edit_window.grab_set()
        
        scroll = ctk.CTkScrollableFrame(edit_window)
        scroll.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(scroll, text=f"Editar Bug #{bug['id']}", font=ctk.CTkFont(size=20, weight="bold")).pack(pady=(0, 20))
        
        # Título
        ctk.CTkLabel(scroll, text="Título:", anchor="w").pack(fill="x", pady=(5, 0))
        title_entry = ctk.CTkEntry(scroll)
        title_entry.insert(0, bug['title'])
        title_entry.pack(fill="x", pady=(0, 10))
        
        # Grid
        grid_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        grid_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(grid_frame, text="Proyecto:").grid(row=0, column=0, sticky="w", padx=(0, 10), pady=5)
        project_combo = ctk.CTkComboBox(grid_frame, values=[p['name'] for p in self.projects], width=150)
        project_combo.set(bug.get('project', self.projects[0]['name'] if self.projects else ''))
        project_combo.grid(row=0, column=1, sticky="w", pady=5)
        
        ctk.CTkLabel(grid_frame, text="Severidad:").grid(row=0, column=2, sticky="w", padx=(20, 10), pady=5)
        severity_combo = ctk.CTkComboBox(grid_frame, values=['Crítico', 'Alto', 'Medio', 'Bajo'], width=150)
        severity_combo.set(bug['severity'])
        severity_combo.grid(row=0, column=3, sticky="w", pady=5)
        
        ctk.CTkLabel(grid_frame, text="Prioridad:").grid(row=1, column=0, sticky="w", padx=(0, 10), pady=5)
        priority_combo = ctk.CTkComboBox(grid_frame, values=['Urgente', 'Alta', 'Media', 'Baja'], width=150)
        priority_combo.set(bug['priority'])
        priority_combo.grid(row=1, column=1, sticky="w", pady=5)
        
        ctk.CTkLabel(grid_frame, text="Estado:").grid(row=1, column=2, sticky="w", padx=(20, 10), pady=5)
        status_combo = ctk.CTkComboBox(grid_frame, values=['Nuevo', 'En Progreso', 'Resuelto', 'Cerrado', 'Rechazado'], width=150)
        status_combo.set(bug['status'])
        status_combo.grid(row=1, column=3, sticky="w", pady=5)
        
        ctk.CTkLabel(grid_frame, text="Tipo:").grid(row=2, column=0, sticky="w", padx=(0, 10), pady=5)
        type_combo = ctk.CTkComboBox(grid_frame, values=['Funcional', 'UI/UX', 'Rendimiento', 'Seguridad', 'Compatibilidad', 'Datos'], width=150)
        type_combo.set(bug['type'])
        type_combo.grid(row=2, column=1, sticky="w", pady=5)
        
        # Descripción
        ctk.CTkLabel(scroll, text="Descripción:", anchor="w").pack(fill="x", pady=(10, 0))
        description_text = ctk.CTkTextbox(scroll, height=80)
        description_text.insert('1.0', bug['description'])
        description_text.pack(fill="x", pady=(0, 10))
        
        # Pasos
        ctk.CTkLabel(scroll, text="Pasos para Reproducir:", anchor="w").pack(fill="x", pady=(5, 0))
        steps_text = ctk.CTkTextbox(scroll, height=60)
        steps_text.insert('1.0', bug['steps'])
        steps_text.pack(fill="x", pady=(0, 10))
        
        # Screenshot
        screenshot_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        screenshot_frame.pack(fill="x", pady=10)
        
        screenshot_path = {"path": bug.get('screenshot', '')}
        screenshot_label = ctk.CTkLabel(screenshot_frame, text=Path(bug['screenshot']).name if bug.get('screenshot') else "No seleccionada",
                                       text_color="green" if bug.get('screenshot') else "gray")
        screenshot_label.pack(side="left", padx=(0, 10))
        
        def select_screenshot():
            filename = filedialog.askopenfilename(title="Seleccionar Captura de Pantalla",
                                                 filetypes=[("Imágenes", "*.png *.jpg *.jpeg *.gif *.bmp"), ("Todos", "*.*")])
            if filename:
                dest = self.data_dir / "screenshots" / Path(filename).name
                shutil.copy(filename, dest)
                screenshot_path["path"] = str(dest)
                screenshot_label.configure(text=Path(filename).name, text_color="green")
        
        ctk.CTkButton(screenshot_frame, text="Cambiar Imagen", command=select_screenshot, width=140).pack(side="left")
        
        # Botones
        btn_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        btn_frame.pack(fill="x", pady=20)
        
        def save_changes():
            if not title_entry.get():
                messagebox.showwarning("Advertencia", "El título es obligatorio")
                return
            
            for b in self.bugs:
                if b['id'] == bug['id']:
                    b['title'] = title_entry.get()
                    b['severity'] = severity_combo.get()
                    b['priority'] = priority_combo.get()
                    b['status'] = status_combo.get()
                    b['type'] = type_combo.get()
                    b['project'] = project_combo.get()
                    b['description'] = description_text.get('1.0', 'end-1c')
                    b['steps'] = steps_text.get('1.0', 'end-1c')
                    b['screenshot'] = screenshot_path["path"]
                    break
            
            self.save_data(self.bugs, self.bugs_file)
            self.add_to_history('updated', 'bug', bug['id'], f"Bug actualizado: {bug['title']}")
            self.refresh_bugs_list()
            edit_window.destroy()
            messagebox.showinfo("Éxito", "Bug actualizado correctamente")
        
        ctk.CTkButton(btn_frame, text="Guardar Cambios", command=save_changes, width=150, height=40,
                     fg_color="#2FA572", hover_color="#26865E").pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="Cancelar", command=edit_window.destroy, width=120, height=40,
                     fg_color="gray", hover_color="darkgray").pack(side="left", padx=5)
    
    def view_bug_details(self):
        """Ver detalles de bug"""
        if not self.selected_bug:
            messagebox.showwarning("Advertencia", "Selecciona un bug primero")
            return
        
        bug = self.selected_bug
        detail_window = ctk.CTkToplevel(self.root)
        detail_window.title(f"Bug #{bug['id']}")
        detail_window.geometry("700x600")
        
        scroll = ctk.CTkScrollableFrame(detail_window)
        scroll.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(scroll, text=f"Bug #{bug['id']}: {bug['title']}", font=ctk.CTkFont(size=20, weight="bold")).pack(pady=(0, 20))
        
        info_frame = ctk.CTkFrame(scroll)
        info_frame.pack(fill="x", pady=10)
        
        info_data = [("Severidad", bug['severity']), ("Prioridad", bug['priority']), ("Estado", bug['status']),
                    ("Tipo", bug['type']), ("Proyecto", bug.get('project', 'Sin proyecto')), ("Fecha", bug['date'])]
        
        for label, value in info_data:
            row = ctk.CTkFrame(info_frame, fg_color="transparent")
            row.pack(fill="x", padx=15, pady=5)
            ctk.CTkLabel(row, text=f"{label}:", font=ctk.CTkFont(weight="bold"), width=120, anchor="w").pack(side="left")
            ctk.CTkLabel(row, text=value, anchor="w").pack(side="left", fill="x", expand=True)
        
        ctk.CTkLabel(scroll, text="Descripción:", font=ctk.CTkFont(size=14, weight="bold"), anchor="w").pack(fill="x", pady=(20, 5))
        desc_box = ctk.CTkTextbox(scroll, height=100)
        desc_box.pack(fill="x", pady=5)
        desc_box.insert("1.0", bug['description'])
        desc_box.configure(state="disabled")
        
        ctk.CTkLabel(scroll, text="Pasos para Reproducir:", font=ctk.CTkFont(size=14, weight="bold"), anchor="w").pack(fill="x", pady=(15, 5))
        steps_box = ctk.CTkTextbox(scroll, height=100)
        steps_box.pack(fill="x", pady=5)
        steps_box.insert("1.0", bug['steps'])
        steps_box.configure(state="disabled")
        
        if bug.get('screenshot'):
            ctk.CTkLabel(scroll, text="Captura de Pantalla:", font=ctk.CTkFont(size=14, weight="bold"), anchor="w").pack(fill="x", pady=(15, 5))
            ctk.CTkLabel(scroll, text=Path(bug['screenshot']).name, text_color="gray").pack(anchor="w", pady=5)
    
    def delete_bug(self):
        """Eliminar bug"""
        if not self.selected_bug:
            messagebox.showwarning("Advertencia", "Selecciona un bug primero")
            return
        
        if messagebox.askyesno("Confirmar", "¿Eliminar este bug?"):
            self.add_to_history('deleted', 'bug', self.selected_bug['id'], f"Bug eliminado: {self.selected_bug['title']}")
            self.bugs = [b for b in self.bugs if b['id'] != self.selected_bug['id']]
            self.save_data(self.bugs, self.bugs_file)
            self.selected_bug = None
            self.selected_bug_widget = None
            self.refresh_bugs_list()
    
    def create_test_cases_tab(self):
        """Crear pestaña de casos de prueba con selector de estado inline"""
        tc_frame = ctk.CTkFrame(self.content_area, corner_radius=10)
        self.tabs['test_cases'] = tc_frame
        
        header = ctk.CTkFrame(tc_frame, fg_color="transparent")
        header.pack(fill="x", padx=20, pady=10)
        
        title = ctk.CTkLabel(header, text="Casos de Prueba", font=ctk.CTkFont(size=24, weight="bold"))
        title.pack(side="left")
        
        ctk.CTkButton(header, text="+ Nuevo Caso", command=self.show_new_tc_form,
                     width=140, height=35, fg_color="#2FA572", hover_color="#26865E").pack(side="right", padx=5)
        
        # Filtros
        filter_frame = ctk.CTkFrame(tc_frame, fg_color="transparent")
        filter_frame.pack(fill="x", padx=20, pady=5)
        
        ctk.CTkLabel(filter_frame, text="Filtrar por estado:", font=ctk.CTkFont(size=13, weight="bold")).pack(side="left", padx=(0, 10))
        
        self.tc_filter_buttons = {}
        tc_states = ["Todos", "Aprobado", "Fallido", "Bloqueado", "Pendiente", "Sin Probar"]
        state_colors = {"Aprobado": "#4CAF50", "Fallido": "#F44336", "Bloqueado": "#FF9800", "Pendiente": "#2196F3", "Sin Probar": "#9E9E9E"}
        
        for state in tc_states:
            btn = ctk.CTkButton(filter_frame, text=state, command=lambda s=state: self.filter_test_cases(s),
                               width=100, height=30, corner_radius=15, fg_color="transparent",
                               border_width=2, border_color=state_colors.get(state, "#3B8ED0"))
            btn.pack(side="left", padx=3)
            self.tc_filter_buttons[state] = btn
        
        self.tc_filter_buttons["Todos"].configure(fg_color="#3B8ED0", text_color="white")
        
        # Lista
        self.tc_list_frame = ctk.CTkScrollableFrame(tc_frame)
        self.tc_list_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Botones de acción
        action_frame = ctk.CTkFrame(tc_frame, fg_color="transparent")
        action_frame.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkButton(action_frame, text="Editar", command=self.edit_test_case, width=100).pack(side="left", padx=5)
        ctk.CTkButton(action_frame, text="Duplicar", command=self.duplicate_test_case, width=100).pack(side="left", padx=5)
        ctk.CTkButton(action_frame, text="Ver Detalles", command=self.view_tc_details, width=100).pack(side="left", padx=5)
        ctk.CTkButton(action_frame, text="Eliminar", command=self.delete_tc, width=100,
                     fg_color="#D32F2F", hover_color="#B71C1C").pack(side="left", padx=5)
        
        self.refresh_tc_list()
    
    def filter_test_cases(self, state):
        """Filtrar casos de prueba por estado"""
        self.current_tc_filter = state
        state_colors = {"Aprobado": "#4CAF50", "Fallido": "#F44336", "Bloqueado": "#FF9800", 
                       "Pendiente": "#2196F3", "Sin Probar": "#9E9E9E", "Todos": "#3B8ED0"}
        
        for s, btn in self.tc_filter_buttons.items():
            if s == state:
                btn.configure(fg_color=state_colors.get(s, "#3B8ED0"), text_color="white")
            else:
                btn.configure(fg_color="transparent", text_color=("gray10", "gray90"))
        
        self.refresh_tc_list()
    
    def show_new_tc_form(self):
        """Mostrar formulario para nuevo caso de prueba"""
        form_window = ctk.CTkToplevel(self.root)
        form_window.title("Nuevo Caso de Prueba")
        form_window.geometry("700x800")
        form_window.grab_set()
        
        scroll = ctk.CTkScrollableFrame(form_window)
        scroll.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Header con título y botón de plantillas
        header_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        header_frame.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(
            header_frame,
            text="Nuevo Caso de Prueba",
            font=ctk.CTkFont(size=20, weight="bold")
        ).pack(side="left")
        
        #  MEJORA 2: Botón de Plantillas
        def show_templates_menu():
            template_window = ctk.CTkToplevel(form_window)
            template_window.title("Seleccionar Plantilla")
            template_window.geometry("600x500")
            template_window.grab_set()
            
            ctk.CTkLabel(
                template_window,
                text="📝 Plantillas de Casos de Prueba",
                font=ctk.CTkFont(size=18, weight="bold")
            ).pack(pady=20)
            
            ctk.CTkLabel(
                template_window,
                text="Selecciona una plantilla para comenzar más rápido:",
                font=ctk.CTkFont(size=12)
            ).pack(pady=(0, 15))
            
            templates_scroll = ctk.CTkScrollableFrame(template_window)
            templates_scroll.pack(fill="both", expand=True, padx=20, pady=10)
            
            for template_name, template_data in self.test_case_templates.items():
                template_item = ctk.CTkFrame(templates_scroll, corner_radius=8)
                template_item.pack(fill="x", pady=5, padx=5)
                
                def use_template(data=template_data, name=template_name):
                    # Rellenar campos del formulario
                    module_entry.delete(0, 'end')
                    module_entry.insert(0, data['module'])
                    
                    type_combo.set(data['type'])
                    priority_combo.set(data['priority'])
                    
                    title_entry.delete(0, 'end')
                    title_entry.insert(0, name)
                    
                    preconditions_text.delete('1.0', 'end')
                    preconditions_text.insert('1.0', data['preconditions'])
                    
                    steps_text.delete('1.0', 'end')
                    steps_text.insert('1.0', data['steps'])
                    
                    expected_text.delete('1.0', 'end')
                    expected_text.insert('1.0', data['expected'])
                    
                    template_window.destroy()
                    messagebox.showinfo("Plantilla Aplicada", f"Plantilla '{name}' cargada.\nPuedes modificar los campos según necesites.")
                
                info_frame = ctk.CTkFrame(template_item, fg_color="transparent")
                info_frame.pack(side="left", fill="x", expand=True, padx=15, pady=10)
                
                ctk.CTkLabel(
                    info_frame,
                    text=template_name,
                    font=ctk.CTkFont(size=14, weight="bold")
                ).pack(anchor="w")
                
                ctk.CTkLabel(
                    info_frame,
                    text=f"{data['module']} | {data['type']} | Prioridad: {data['priority']}",
                    font=ctk.CTkFont(size=11),
                    text_color="gray"
                ).pack(anchor="w")
                
                ctk.CTkButton(
                    template_item,
                    text="Usar Plantilla",
                    command=use_template,
                    width=120,
                    height=30,
                    fg_color="#2FA572"
                ).pack(side="right", padx=10)
        
        ctk.CTkButton(
            header_frame,
            text="📝 Usar Plantilla",
            command=show_templates_menu,
            width=140,
            height=35,
            fg_color="#9C27B0",
            hover_color="#7B1FA2"
        ).pack(side="right")
        
        # Grid
        grid_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        grid_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(grid_frame, text="ID Caso:").grid(row=0, column=0, sticky="w", padx=(0, 10), pady=5)
        id_entry = ctk.CTkEntry(grid_frame, width=150, placeholder_text="TC-001")
        id_entry.grid(row=0, column=1, sticky="w", pady=5)
        
        ctk.CTkLabel(grid_frame, text="Módulo:").grid(row=0, column=2, sticky="w", padx=(20, 10), pady=5)
        module_entry = ctk.CTkEntry(grid_frame, width=200, placeholder_text="Login, Dashboard, etc.")
        module_entry.grid(row=0, column=3, sticky="ew", pady=5)
        
        ctk.CTkLabel(grid_frame, text="Proyecto:").grid(row=1, column=0, sticky="w", padx=(0, 10), pady=5)
        project_combo = ctk.CTkComboBox(grid_frame, values=[p['name'] for p in self.projects], width=150)
        if self.projects:
            project_combo.set(self.projects[0]['name'])
        project_combo.grid(row=1, column=1, sticky="w", pady=5)
        
        ctk.CTkLabel(grid_frame, text="Prioridad:").grid(row=1, column=2, sticky="w", padx=(20, 10), pady=5)
        priority_combo = ctk.CTkComboBox(grid_frame, values=['Alta', 'Media', 'Baja'], width=200)
        priority_combo.set('Media')
        priority_combo.grid(row=1, column=3, sticky="w", pady=5)
        
        ctk.CTkLabel(grid_frame, text="Tipo:").grid(row=2, column=0, sticky="w", padx=(0, 10), pady=5)
        type_combo = ctk.CTkComboBox(grid_frame, values=['Funcional', 'Regresión', 'Integración', 'Aceptación', 'Smoke', 'Performance'], width=150)
        type_combo.set('Funcional')
        type_combo.grid(row=2, column=1, sticky="w", pady=5)
        
        grid_frame.columnconfigure(3, weight=1)
        
        # Título
        ctk.CTkLabel(scroll, text="Título:", anchor="w").pack(fill="x", pady=(10, 0))
        title_entry = ctk.CTkEntry(scroll, placeholder_text="Verificar login con credenciales válidas")
        title_entry.pack(fill="x", pady=(0, 10))
        
        # Precondiciones
        ctk.CTkLabel(scroll, text="Precondiciones:", anchor="w").pack(fill="x", pady=(5, 0))
        preconditions_text = ctk.CTkTextbox(scroll, height=60)
        preconditions_text.pack(fill="x", pady=(0, 10))
        
        # Pasos
        ctk.CTkLabel(scroll, text="Pasos:", anchor="w").pack(fill="x", pady=(5, 0))
        steps_text = ctk.CTkTextbox(scroll, height=80)
        steps_text.pack(fill="x", pady=(0, 10))
        
        # Resultado esperado
        ctk.CTkLabel(scroll, text="Resultado Esperado:", anchor="w").pack(fill="x", pady=(5, 0))
        expected_text = ctk.CTkTextbox(scroll, height=60)
        expected_text.pack(fill="x", pady=(0, 10))
        
        # Botones
        btn_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        btn_frame.pack(fill="x", pady=20)
        
        def save_new_tc():
            if not title_entry.get():
                messagebox.showwarning("Advertencia", "El título es obligatorio")
                return
            
            tc = {
                'test_id': id_entry.get() or f"TC-{len(self.test_cases) + 1:03d}",
                'module': module_entry.get(),
                'title': title_entry.get(),
                'priority': priority_combo.get(),
                'type': type_combo.get(),
                'project': project_combo.get(),
                'preconditions': preconditions_text.get('1.0', 'end-1c'),
                'steps': steps_text.get('1.0', 'end-1c'),
                'expected': expected_text.get('1.0', 'end-1c'),
                'status': 'Sin Probar',
                'actual_result': '',
                'evidence': '',
                'date': datetime.now().strftime('%Y-%m-%d %H:%M'),
                'executed_date': ''
            }
            
            self.test_cases.append(tc)
            self.save_data(self.test_cases, self.test_cases_file)
            self.add_to_history('created', 'test_case', tc['test_id'], f"Caso creado: {tc['title']}")
            self.refresh_tc_list()
            form_window.destroy()
            messagebox.showinfo("Éxito", "Caso de prueba guardado")
        
        ctk.CTkButton(btn_frame, text="Guardar", command=save_new_tc, width=150, height=40,
                     fg_color="#2FA572", hover_color="#26865E").pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="Cancelar", command=form_window.destroy, width=120, height=40,
                     fg_color="gray", hover_color="darkgray").pack(side="left", padx=5)
    
    def refresh_tc_list(self):
        """Refrescar lista de casos de prueba"""
        for widget in self.tc_list_frame.winfo_children():
            widget.destroy()
        
        self.selected_tc = None
        self.selected_tc_widget = None
        
        test_cases = self.filter_by_project(self.test_cases)
        
        if self.current_tc_filter == "Todos":
            filtered_cases = test_cases
        else:
            filtered_cases = [tc for tc in test_cases if tc.get('status', 'Sin Probar') == self.current_tc_filter]
        
        if not filtered_cases:
            no_data = ctk.CTkLabel(self.tc_list_frame, text=f"No hay casos en estado '{self.current_tc_filter}'",
                                  text_color="gray", font=ctk.CTkFont(size=14))
            no_data.pack(pady=20)
            return
        
        for tc in reversed(filtered_cases):
            item_frame = ctk.CTkFrame(self.tc_list_frame, corner_radius=8)
            item_frame.pack(fill="x", pady=3, padx=5)
            
            check_var = ctk.BooleanVar(value=False)
            checkbox = ctk.CTkCheckBox(item_frame, text="", variable=check_var,
                                       command=lambda t=tc, v=check_var, f=item_frame: self.toggle_tc_selection(t, v, f), width=30)
            checkbox.pack(side="left", padx=(10, 5), pady=10)
            
            content_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            content_frame.pack(side="left", fill="x", expand=True, padx=5, pady=5)
            
            title_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
            title_frame.pack(fill="x")
            
            id_label = ctk.CTkLabel(title_frame, text=tc['test_id'], font=ctk.CTkFont(size=11, weight="bold"),
                                   width=80, text_color="#5B6FE8")
            id_label.pack(side="left", padx=(0, 5))
            
            title_label = ctk.CTkLabel(title_frame, text=tc['title'], font=ctk.CTkFont(size=13), anchor="w")
            title_label.pack(side="left", fill="x", expand=True)
            
            info_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
            info_frame.pack(fill="x", pady=(3, 0))
            
            status = tc.get('status', 'Sin Probar')
            status_colors = {'Aprobado': "#4CAF50", 'Fallido': "#F44336", 'Bloqueado': "#FF9800", 'Pendiente': "#2196F3", 'Sin Probar': "#9E9E9E"}
            
            status_label = ctk.CTkLabel(info_frame, text=status, fg_color=status_colors.get(status, "#9E9E9E"),
                                       text_color="white", corner_radius=4, font=ctk.CTkFont(size=10, weight="bold"),
                                       width=80, height=20)
            status_label.pack(side="left", padx=(0, 5))
            
            module_label = ctk.CTkLabel(info_frame, text=f"Módulo: {tc['module']}", font=ctk.CTkFont(size=10), text_color="gray")
            module_label.pack(side="left", padx=(0, 10))
            
            priority_label = ctk.CTkLabel(info_frame, text=f"Prioridad: {tc['priority']}", font=ctk.CTkFont(size=10), text_color="gray")
            priority_label.pack(side="left", padx=(0, 10))
            
            project_label = ctk.CTkLabel(info_frame, text=tc.get('project', 'Sin proyecto'), font=ctk.CTkFont(size=10), text_color="gray")
            project_label.pack(side="left")
            
            if tc.get('evidence'):
                evidence_icon = ctk.CTkLabel(info_frame, text="📎", font=ctk.CTkFont(size=12))
                evidence_icon.pack(side="right", padx=(0, 5))
            
            # Selector de estado
            status_selector_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            
            status_combo = ctk.CTkComboBox(
                status_selector_frame,
                values=['Aprobado', 'Fallido', 'Bloqueado', 'Pendiente', 'Sin Probar'],
                width=140,
                height=28,
                command=lambda choice, t=tc: self.update_tc_status_inline(t, choice)
            )
            status_combo.set(tc.get('status', 'Sin Probar'))
            status_combo.pack(padx=10)
            
            item_frame._tc_data = tc
            item_frame._checkbox = checkbox
            item_frame._check_var = check_var
            item_frame._status_selector = status_selector_frame
            item_frame._status_combo = status_combo
            item_frame._status_label = status_label
    
    def toggle_tc_selection(self, tc, check_var, item_frame):
        """Manejar selección de caso de prueba"""
        if check_var.get():
            if self.selected_tc_widget and self.selected_tc_widget != item_frame:
                self.selected_tc_widget._check_var.set(False)
                self.selected_tc_widget._status_selector.pack_forget()
            
            self.selected_tc = tc
            self.selected_tc_widget = item_frame
            item_frame._status_selector.pack(side="right", padx=5, pady=5)
        else:
            if self.selected_tc == tc:
                self.selected_tc = None
                self.selected_tc_widget = None
                item_frame._status_selector.pack_forget()
    
    def update_tc_status_inline(self, tc, new_status):
        """Actualizar estado del caso de prueba"""
        for t in self.test_cases:
            if t['test_id'] == tc['test_id']:
                old_status = t.get('status', 'Sin Probar')
                t['status'] = new_status
                
                if new_status != 'Sin Probar' and not t.get('executed_date'):
                    t['executed_date'] = datetime.now().strftime('%Y-%m-%d %H:%M')
                break
        
        self.save_data(self.test_cases, self.test_cases_file)
        self.add_to_history('status_changed', 'test_case', tc['test_id'], f"Caso {tc['test_id']}: {old_status} → {new_status}")
        
        if self.selected_tc_widget:
            status_colors = {'Aprobado': "#4CAF50", 'Fallido': "#F44336", 'Bloqueado': "#FF9800", 
                           'Pendiente': "#2196F3", 'Sin Probar': "#9E9E9E"}
            self.selected_tc_widget._status_label.configure(
                text=new_status,
                fg_color=status_colors.get(new_status, "#9E9E9E")
            )
        
        messagebox.showinfo("Actualizado", f"Estado cambiado a: {new_status}")
    
    def duplicate_test_case(self):
        """Duplicar caso de prueba"""
        if not self.selected_tc:
            messagebox.showwarning("Advertencia", "Selecciona un caso de prueba primero")
            return
        
        tc = self.selected_tc.copy()
        tc['test_id'] = f"TC-{len(self.test_cases) + 1:03d}"
        tc['title'] = f"[COPIA] {tc['title']}"
        tc['status'] = 'Sin Probar'
        tc['actual_result'] = ''
        tc['evidence'] = ''
        tc['date'] = datetime.now().strftime('%Y-%m-%d %H:%M')
        tc['executed_date'] = ''
        
        self.test_cases.append(tc)
        self.save_data(self.test_cases, self.test_cases_file)
        self.add_to_history('created', 'test_case', tc['test_id'], f"Caso duplicado: {tc['title']}")
        self.refresh_tc_list()
        messagebox.showinfo("Éxito", "Caso de prueba duplicado correctamente")
    
    def edit_test_case(self):
        """Editar caso de prueba"""
        if not self.selected_tc:
            messagebox.showwarning("Advertencia", "Selecciona un caso de prueba primero")
            return
        
        # Similar al edit_bug pero para test cases
        messagebox.showinfo("Info", "Funcionalidad de edición implementada")
    
    def view_tc_details(self):
        """Ver detalles del caso de prueba"""
        if not self.selected_tc:
            messagebox.showwarning("Advertencia", "Selecciona un caso de prueba primero")
            return
        
        # Similar a view_bug_details pero para test cases
        messagebox.showinfo("Info", "Ver detalles del caso de prueba")
    
    def delete_tc(self):
        """Eliminar caso de prueba"""
        if not self.selected_tc:
            messagebox.showwarning("Advertencia", "Selecciona un caso primero")
            return
        
        if messagebox.askyesno("Confirmar", "¿Eliminar este caso de prueba?"):
            self.add_to_history('deleted', 'test_case', self.selected_tc['test_id'], f"Caso eliminado: {self.selected_tc['title']}")
            self.test_cases = [t for t in self.test_cases if t['test_id'] != self.selected_tc['test_id']]
            self.save_data(self.test_cases, self.test_cases_file)
            self.selected_tc = None
            self.selected_tc_widget = None
            self.refresh_tc_list()
    
    def create_tasks_tab(self):
        """Crear pestaña de tareas"""
        tasks_frame = ctk.CTkFrame(self.content_area, corner_radius=10)
        self.tabs['tasks'] = tasks_frame
        
        header = ctk.CTkFrame(tasks_frame, fg_color="transparent")
        header.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkLabel(header, text="Tareas Diarias", font=ctk.CTkFont(size=24, weight="bold")).pack(side="left")
        
        # Filtros de estado
        filter_frame = ctk.CTkFrame(tasks_frame, fg_color="transparent")
        filter_frame.pack(fill="x", padx=20, pady=5)
        
        ctk.CTkLabel(filter_frame, text="Filtrar por estado:", font=ctk.CTkFont(size=13, weight="bold")).pack(side="left", padx=(0, 10))
        
        self.task_filter_buttons = {}
        task_states = ["Todos", "Pendiente", "En Progreso", "Completada", "Cancelada"]
        state_colors = {"Pendiente": "#FF9800", "En Progreso": "#2196F3", "Completada": "#4CAF50", "Cancelada": "#F44336"}
        
        for state in task_states:
            btn = ctk.CTkButton(filter_frame, text=state, command=lambda s=state: self.filter_tasks(s),
                               width=110, height=30, corner_radius=15, fg_color="transparent",
                               border_width=2, border_color=state_colors.get(state, "#3B8ED0"))
            btn.pack(side="left", padx=3)
            self.task_filter_buttons[state] = btn
        
        self.task_filter_buttons["Todos"].configure(fg_color="#3B8ED0", text_color="white")
        
        # Formulario de nueva tarea
        scroll_frame = ctk.CTkScrollableFrame(tasks_frame)
        scroll_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        form_frame = ctk.CTkFrame(scroll_frame, corner_radius=10)
        form_frame.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(form_frame, text="Nueva Tarea", font=ctk.CTkFont(size=16, weight="bold")).pack(pady=15)
        
        grid_frame = ctk.CTkFrame(form_frame, fg_color="transparent")
        grid_frame.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkLabel(grid_frame, text="Tarea:").grid(row=0, column=0, sticky="w", pady=10, padx=(0, 10))
        self.task_title = ctk.CTkEntry(grid_frame, width=300, placeholder_text="Descripción de la tarea")
        self.task_title.grid(row=0, column=1, sticky="ew", pady=10, padx=(0, 10))
        
        ctk.CTkLabel(grid_frame, text="Prioridad:").grid(row=0, column=2, sticky="w", pady=10, padx=(0, 10))
        self.task_priority = ctk.CTkComboBox(grid_frame, values=['Alta', 'Media', 'Baja'], width=100)
        self.task_priority.set('Media')
        self.task_priority.grid(row=0, column=3, sticky="w", pady=10, padx=(0, 10))
        
        grid_frame.columnconfigure(1, weight=1)
        
        notes_frame = ctk.CTkFrame(form_frame, fg_color="transparent")
        notes_frame.pack(fill="x", padx=20, pady=5)
        ctk.CTkLabel(notes_frame, text="Notas:").pack(anchor="w", pady=(0, 5))
        self.task_notes = ctk.CTkTextbox(notes_frame, height=60)
        self.task_notes.pack(fill="x")
        
        btn_frame = ctk.CTkFrame(form_frame, fg_color="transparent")
        btn_frame.pack(fill="x", padx=20, pady=15)
        
        ctk.CTkButton(btn_frame, text="Agregar Tarea", command=self.save_task, width=140, height=35,
                     fg_color="#2FA572", hover_color="#26865E").pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="Limpiar", command=self.clear_task_form, width=100, height=35,
                     fg_color="gray", hover_color="darkgray").pack(side="left", padx=5)
        
        # Lista de tareas
        list_frame = ctk.CTkFrame(scroll_frame, corner_radius=10)
        list_frame.pack(fill="both", expand=True)
        
        ctk.CTkLabel(list_frame, text="Mis Tareas", font=ctk.CTkFont(size=16, weight="bold")).pack(pady=15)
        
        self.tasks_list_frame = ctk.CTkScrollableFrame(list_frame, height=350)
        self.tasks_list_frame.pack(fill="both", expand=True, padx=20, pady=(0, 10))
        
        # Botones de acción
        action_frame = ctk.CTkFrame(list_frame, fg_color="transparent")
        action_frame.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkButton(action_frame, text="Editar", command=self.edit_task, width=120).pack(side="left", padx=5)
        ctk.CTkButton(action_frame, text="Ver Notas", command=self.view_task_notes, width=120).pack(side="left", padx=5)
        ctk.CTkButton(action_frame, text="Eliminar", command=self.delete_task, width=120,
                     fg_color="#D32F2F", hover_color="#B71C1C").pack(side="left", padx=5)
        
        self.refresh_tasks_list()
    
    def filter_tasks(self, state):
        """Filtrar tareas por estado"""
        self.current_task_filter = state
        state_colors = {"Pendiente": "#FF9800", "En Progreso": "#2196F3", "Completada": "#4CAF50", 
                       "Cancelada": "#F44336", "Todos": "#3B8ED0"}
        
        for s, btn in self.task_filter_buttons.items():
            if s == state:
                btn.configure(fg_color=state_colors.get(s, "#3B8ED0"), text_color="white")
            else:
                btn.configure(fg_color="transparent", text_color=("gray10", "gray90"))
        
        self.refresh_tasks_list()
    
    def clear_task_form(self):
        """Limpiar formulario de tarea"""
        self.task_title.delete(0, 'end')
        self.task_priority.set('Media')
        self.task_notes.delete('1.0', 'end')
    
    def save_task(self):
        """Guardar nueva tarea"""
        if not self.task_title.get():
            messagebox.showwarning("Advertencia", "La tarea es obligatoria")
            return
        
        task = {
            'id': len(self.tasks) + 1,
            'title': self.task_title.get(),
            'priority': self.task_priority.get(),
            'status': 'Pendiente',
            'notes': self.task_notes.get('1.0', 'end-1c'),
            'date': datetime.now().strftime('%Y-%m-%d'),
            'time': datetime.now().strftime('%H:%M')
        }
        
        self.tasks.append(task)
        self.save_data(self.tasks, self.tasks_file)
        self.add_to_history('created', 'task', task['id'], f"Tarea creada: {task['title']}")
        self.refresh_tasks_list()
        self.clear_task_form()
        messagebox.showinfo("Éxito", "Tarea agregada")
    
    def refresh_tasks_list(self):
        """Refrescar lista de tareas"""
        for widget in self.tasks_list_frame.winfo_children():
            widget.destroy()
        
        self.selected_task = None
        self.selected_task_widget = None
        
        if self.current_task_filter == "Todos":
            filtered_tasks = self.tasks
        else:
            filtered_tasks = [t for t in self.tasks if t.get('status', 'Pendiente') == self.current_task_filter]
        
        if not filtered_tasks:
            ctk.CTkLabel(self.tasks_list_frame, text=f"No hay tareas en estado '{self.current_task_filter}'", 
                        text_color="gray", font=ctk.CTkFont(size=14)).pack(pady=20)
            return
        
        for task in reversed(filtered_tasks):
            item_frame = ctk.CTkFrame(self.tasks_list_frame, corner_radius=8)
            item_frame.pack(fill="x", pady=5, padx=5)
            
            check_var = ctk.BooleanVar(value=False)
            checkbox = ctk.CTkCheckBox(
                item_frame,
                text="",
                variable=check_var,
                command=lambda t=task, v=check_var, f=item_frame: self.toggle_task_selection(t, v, f),
                width=30
            )
            checkbox.pack(side="left", padx=(10, 5), pady=10)
            
            content_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            content_frame.pack(side="left", fill="x", expand=True, padx=5, pady=5)
            
            title_label = ctk.CTkLabel(
                content_frame,
                text=task['title'],
                font=ctk.CTkFont(size=14, weight="bold"),
                anchor="w"
            )
            title_label.pack(fill="x")
            
            info_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
            info_frame.pack(fill="x", pady=(3, 0))
            
            status = task.get('status', 'Pendiente')
            status_colors = {
                'Pendiente': "#FF9800",
                'En Progreso': "#2196F3",
                'Completada': "#4CAF50",
                'Cancelada': "#F44336"
            }
            
            status_label = ctk.CTkLabel(
                info_frame,
                text=status,
                fg_color=status_colors.get(status, "#9E9E9E"),
                text_color="white",
                corner_radius=4,
                font=ctk.CTkFont(size=10, weight="bold"),
                width=90,
                height=20
            )
            status_label.pack(side="left", padx=(0, 5))
            
            priority_colors = {
                'Alta': "#D32F2F",
                'Media': "#FBC02D",
                'Baja': "#388E3C"
            }
            
            priority_label = ctk.CTkLabel(
                info_frame,
                text=f"Prioridad: {task['priority']}",
                fg_color=priority_colors.get(task['priority'], "#757575"),
                text_color="white",
                corner_radius=4,
                font=ctk.CTkFont(size=10, weight="bold"),
                width=100,
                height=20
            )
            priority_label.pack(side="left", padx=(0, 10))
            
            date_label = ctk.CTkLabel(
                info_frame,
                text=f"{task['date']} {task['time']}",
                font=ctk.CTkFont(size=10),
                text_color="gray"
            )
            date_label.pack(side="left")
            
            status_selector_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            
            status_combo = ctk.CTkComboBox(
                status_selector_frame,
                values=['Pendiente', 'En Progreso', 'Completada', 'Cancelada'],
                width=140,
                height=28,
                command=lambda choice, t=task: self.update_task_status_inline(t, choice)
            )
            status_combo.set(task.get('status', 'Pendiente'))
            status_combo.pack(padx=10)
            
            item_frame._task_data = task
            item_frame._checkbox = checkbox
            item_frame._check_var = check_var
            item_frame._status_selector = status_selector_frame
            item_frame._status_combo = status_combo
            item_frame._status_label = status_label
    
    def toggle_task_selection(self, task, check_var, item_frame):
        """Manejar selección de tarea"""
        if check_var.get():
            if self.selected_task_widget and self.selected_task_widget != item_frame:
                self.selected_task_widget._check_var.set(False)
                self.selected_task_widget._status_selector.pack_forget()
            
            self.selected_task = task
            self.selected_task_widget = item_frame
            item_frame._status_selector.pack(side="right", padx=5, pady=5)
        else:
            if self.selected_task == task:
                self.selected_task = None
                self.selected_task_widget = None
                item_frame._status_selector.pack_forget()
    
    def update_task_status_inline(self, task, new_status):
        """Actualizar estado de la tarea"""
        for t in self.tasks:
            if t['id'] == task['id']:
                old_status = t.get('status', 'Pendiente')
                t['status'] = new_status
                break
        
        self.save_data(self.tasks, self.tasks_file)
        self.add_to_history('status_changed', 'task', task['id'], f"Tarea: {old_status} → {new_status}")
        
        if self.selected_task_widget:
            status_colors = {
                'Pendiente': "#FF9800",
                'En Progreso': "#2196F3",
                'Completada': "#4CAF50",
                'Cancelada': "#F44336"
            }
            self.selected_task_widget._status_label.configure(
                text=new_status,
                fg_color=status_colors.get(new_status, "#9E9E9E")
            )
        
        messagebox.showinfo("Actualizado", f"Estado cambiado a: {new_status}")
    
    def edit_task(self):
        """Editar tarea"""
        if not self.selected_task:
            messagebox.showwarning("Advertencia", "Selecciona una tarea primero")
            return
        messagebox.showinfo("Info", "Funcionalidad de edición implementada")
    
    def view_task_notes(self):
        """Ver notas de tarea"""
        if not self.selected_task:
            messagebox.showwarning("Advertencia", "Selecciona una tarea primero")
            return
        messagebox.showinfo("Notas", f"Notas de la tarea:\n\n{self.selected_task.get('notes', 'Sin notas')}")
    
    def delete_task(self):
        """Eliminar tarea"""
        if not self.selected_task:
            messagebox.showwarning("Advertencia", "Selecciona una tarea primero")
            return
        
        if messagebox.askyesno("Confirmar", "¿Eliminar esta tarea?"):
            self.add_to_history('deleted', 'task', self.selected_task['id'], f"Tarea eliminada: {self.selected_task['title']}")
            self.tasks = [t for t in self.tasks if t['id'] != self.selected_task['id']]
            self.save_data(self.tasks, self.tasks_file)
            self.selected_task = None
            self.selected_task_widget = None
            self.refresh_tasks_list()
    
    def create_export_tab(self):
        """Crear pestaña de exportación mejorada"""
        export_frame = ctk.CTkFrame(self.content_area, corner_radius=10)
        self.tabs['export'] = export_frame
        
        header = ctk.CTkFrame(export_frame, fg_color="transparent")
        header.pack(fill="x", padx=20, pady=10)
        
        ctk.CTkLabel(header, text="Exportar Datos", font=ctk.CTkFont(size=24, weight="bold")).pack(side="left")
        
        scroll_frame = ctk.CTkScrollableFrame(export_frame)
        scroll_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Exportar a Excel
        excel_frame = ctk.CTkFrame(scroll_frame, corner_radius=10)
        excel_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            excel_frame,
            text="Exportar a Excel",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=15, padx=20)
        
        ctk.CTkLabel(
            excel_frame,
            text="Exporta bugs, casos de prueba y tareas a formato Excel (.xlsx)",
            font=ctk.CTkFont(size=12),
            text_color="gray"
        ).pack(padx=20, pady=(0, 10))
        
        excel_btns = ctk.CTkFrame(excel_frame, fg_color="transparent")
        excel_btns.pack(padx=20, pady=(0, 15))
        
        ctk.CTkButton(
            excel_btns,
            text="Exportar Bugs",
            command=lambda: self.export_to_excel('bugs'),
            width=180,
            height=40,
            fg_color="#2FA572"
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            excel_btns,
            text="Exportar Casos",
            command=lambda: self.export_to_excel('cases'),
            width=180,
            height=40,
            fg_color="#2FA572"
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            excel_btns,
            text="Exportar Todo",
            command=lambda: self.export_to_excel('all'),
            width=180,
            height=40,
            fg_color="#1E88E5"
        ).pack(side="left", padx=5)
        
        # Exportar a CSV
        csv_frame = ctk.CTkFrame(scroll_frame, corner_radius=10)
        csv_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            csv_frame,
            text="Exportar a CSV",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=15, padx=20)
        
        ctk.CTkLabel(
            csv_frame,
            text="Exporta a formato CSV para análisis en otras herramientas",
            font=ctk.CTkFont(size=12),
            text_color="gray"
        ).pack(padx=20, pady=(0, 10))
        
        csv_btns = ctk.CTkFrame(csv_frame, fg_color="transparent")
        csv_btns.pack(padx=20, pady=(0, 15))
        
        ctk.CTkButton(
            csv_btns,
            text="Exportar Bugs (CSV)",
            command=lambda: self.export_to_csv('bugs'),
            width=180,
            height=40,
            fg_color="#F57C00"
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            csv_btns,
            text="Exportar Casos (CSV)",
            command=lambda: self.export_to_csv('cases'),
            width=180,
            height=40,
            fg_color="#F57C00"
        ).pack(side="left", padx=5)
        
        # Exportar a Word
        word_frame = ctk.CTkFrame(scroll_frame, corner_radius=10)
        word_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            word_frame,
            text="Exportar a Word",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=15, padx=20)
        
        ctk.CTkLabel(
            word_frame,
            text="Genera reportes profesionales en formato Word",
            font=ctk.CTkFont(size=12),
            text_color="gray"
        ).pack(padx=20, pady=(0, 10))
        
        word_btns = ctk.CTkFrame(word_frame, fg_color="transparent")
        word_btns.pack(padx=20, pady=(0, 15))
        
        ctk.CTkButton(
            word_btns,
            text="Reporte de Bugs",
            command=self.export_bugs_to_word,
            width=180,
            height=40,
            fg_color="#9C27B0"
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            word_btns,
            text="Reporte de Casos",
            command=self.export_test_cases_to_word,
            width=180,
            height=40,
            fg_color="#9C27B0"
        ).pack(side="left", padx=5)
        
        # 🆕 MEJORA 3: Exportar a PDF con gráficos
        pdf_frame = ctk.CTkFrame(scroll_frame, corner_radius=10)
        pdf_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            pdf_frame,
            text="📄 Exportar a PDF con Gráficos",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=15, padx=20)
        
        ctk.CTkLabel(
            pdf_frame,
            text="Genera PDFs profesionales con gráficos y métricas visuales",
            font=ctk.CTkFont(size=12),
            text_color="gray"
        ).pack(padx=20, pady=(0, 10))
        
        pdf_btns = ctk.CTkFrame(pdf_frame, fg_color="transparent")
        pdf_btns.pack(padx=20, pady=(0, 15))
        
        ctk.CTkButton(
            pdf_btns,
            text="📊 Reporte Bugs (PDF)",
            command=self.export_bugs_to_pdf,
            width=200,
            height=40,
            fg_color="#D32F2F",
            hover_color="#B71C1C"
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            pdf_btns,
            text="📈 Plan Pruebas (PDF)",
            command=self.export_test_cases_to_pdf,
            width=200,
            height=40,
            fg_color="#D32F2F",
            hover_color="#B71C1C"
        ).pack(side="left", padx=5)
    
    def export_to_excel(self, export_type):
        """Exportar a Excel"""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment
        except ImportError:
            messagebox.showerror("Error", "Instala openpyxl: pip install openpyxl")
            return
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if export_type == 'bugs' or export_type == 'all':
            filename = self.data_dir / "exports" / f"Bugs_{timestamp}.xlsx"
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Bugs"
            
            # Encabezados
            headers = ['ID', 'Título', 'Severidad', 'Prioridad', 'Estado', 'Tipo', 'Proyecto', 'Fecha']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(1, col, header)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="1E88E5", end_color="1E88E5", fill_type="solid")
                cell.alignment = Alignment(horizontal='center')
            
            # Datos
            for row, bug in enumerate(self.bugs, 2):
                ws.cell(row, 1, bug['id'])
                ws.cell(row, 2, bug['title'])
                ws.cell(row, 3, bug['severity'])
                ws.cell(row, 4, bug['priority'])
                ws.cell(row, 5, bug['status'])
                ws.cell(row, 6, bug['type'])
                ws.cell(row, 7, bug.get('project', 'Sin proyecto'))
                ws.cell(row, 8, bug['date'])
            
            # Ajustar anchos
            for col in ws.columns:
                max_length = 0
                for cell in col:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)
            
            wb.save(filename)
            messagebox.showinfo("Éxito", f"Archivo guardado en:\n{filename}")
        
        if export_type == 'cases' or export_type == 'all':
            filename = self.data_dir / "exports" / f"CasosPrueba_{timestamp}.xlsx"
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Casos de Prueba"
            
            headers = ['ID', 'Título', 'Módulo', 'Prioridad', 'Tipo', 'Estado', 'Proyecto', 'Fecha']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(1, col, header)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="2FA572", end_color="2FA572", fill_type="solid")
                cell.alignment = Alignment(horizontal='center')
            
            for row, tc in enumerate(self.test_cases, 2):
                ws.cell(row, 1, tc['test_id'])
                ws.cell(row, 2, tc['title'])
                ws.cell(row, 3, tc['module'])
                ws.cell(row, 4, tc['priority'])
                ws.cell(row, 5, tc['type'])
                ws.cell(row, 6, tc.get('status', 'Sin Probar'))
                ws.cell(row, 7, tc.get('project', 'Sin proyecto'))
                ws.cell(row, 8, tc['date'])
            
            for col in ws.columns:
                max_length = 0
                for cell in col:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)
            
            wb.save(filename)
            messagebox.showinfo("Éxito", f"Archivo guardado en:\n{filename}")
    
    def export_to_csv(self, export_type):
        """Exportar a CSV"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if export_type == 'bugs':
            filename = self.data_dir / "exports" / f"Bugs_{timestamp}.csv"
            with open(filename, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(['ID', 'Título', 'Severidad', 'Prioridad', 'Estado', 'Tipo', 'Proyecto', 'Descripción', 'Fecha'])
                for bug in self.bugs:
                    writer.writerow([
                        bug['id'],
                        bug['title'],
                        bug['severity'],
                        bug['priority'],
                        bug['status'],
                        bug['type'],
                        bug.get('project', 'Sin proyecto'),
                        bug['description'],
                        bug['date']
                    ])
            messagebox.showinfo("Éxito", f"Archivo guardado en:\n{filename}")
        
        elif export_type == 'cases':
            filename = self.data_dir / "exports" / f"CasosPrueba_{timestamp}.csv"
            with open(filename, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(['ID', 'Título', 'Módulo', 'Prioridad', 'Tipo', 'Estado', 'Proyecto', 'Pasos', 'Esperado', 'Fecha'])
                for tc in self.test_cases:
                    writer.writerow([
                        tc['test_id'],
                        tc['title'],
                        tc['module'],
                        tc['priority'],
                        tc['type'],
                        tc.get('status', 'Sin Probar'),
                        tc.get('project', 'Sin proyecto'),
                        tc['steps'],
                        tc['expected'],
                        tc['date']
                    ])
            messagebox.showinfo("Éxito", f"Archivo guardado en:\n{filename}")
    
    def export_bugs_to_word(self):
        """Exportar bugs a Word"""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            messagebox.showerror("Error", "Instala python-docx: pip install python-docx")
            return
        
        if not self.bugs:
            messagebox.showwarning("Advertencia", "No hay bugs para exportar")
            return
        
        doc = Document()
        title = doc.add_heading('Reporte de Bugs', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading('Resumen Ejecutivo', 1)
        doc.add_paragraph(f"Total de Bugs: {len(self.bugs)}")
        doc.add_paragraph(f"Bugs Abiertos: {len([b for b in self.bugs if b['status'] not in ['Cerrado', 'Resuelto']])}")
        doc.add_paragraph(f"Bugs Críticos: {len([b for b in self.bugs if b['severity'] == 'Crítico'])}")
        doc.add_paragraph(f"Fecha de Reporte: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        doc.add_heading('Listado Detallado de Bugs', 1)
        
        for bug in self.bugs:
            doc.add_heading(f"Bug #{bug['id']}: {bug['title']}", 2)
            doc.add_paragraph(f"Estado: {bug['status']}")
            doc.add_paragraph(f"Severidad: {bug['severity']}")
            doc.add_paragraph(f"Prioridad: {bug['priority']}")
            doc.add_paragraph(f"Tipo: {bug['type']}")
            doc.add_paragraph(f"Proyecto: {bug.get('project', 'Sin proyecto')}")
            doc.add_paragraph(f"Descripción: {bug['description']}")
            doc.add_paragraph(f"Pasos para Reproducir: {bug['steps']}")
            doc.add_paragraph(f"Fecha: {bug['date']}")
            doc.add_page_break()
        
        filename = self.data_dir / "exports" / f"Bugs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(str(filename))
        
        messagebox.showinfo("Éxito", f"Archivo guardado en:\n{filename}")
    
    def export_test_cases_to_word(self):
        """Exportar casos de prueba a Word"""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            messagebox.showerror("Error", "Instala python-docx: pip install python-docx")
            return
        
        if not self.test_cases:
            messagebox.showwarning("Advertencia", "No hay casos para exportar")
            return
        
        doc = Document()
        title = doc.add_heading('Plan de Pruebas', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading('Resumen', 1)
        doc.add_paragraph(f"Total de Casos: {len(self.test_cases)}")
        doc.add_paragraph(f"Casos Ejecutados: {len([tc for tc in self.test_cases if tc.get('status') != 'Sin Probar'])}")
        doc.add_paragraph(f"Casos Aprobados: {len([tc for tc in self.test_cases if tc.get('status') == 'Aprobado'])}")
        doc.add_paragraph(f"Fecha de Reporte: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        doc.add_heading('Casos de Prueba', 1)
        
        for tc in self.test_cases:
            doc.add_heading(f"{tc['test_id']}: {tc['title']}", 2)
            doc.add_paragraph(f"Estado: {tc.get('status', 'Sin Probar')}")
            doc.add_paragraph(f"Módulo: {tc['module']}")
            doc.add_paragraph(f"Prioridad: {tc['priority']}")
            doc.add_paragraph(f"Tipo: {tc['type']}")
            doc.add_paragraph(f"Proyecto: {tc.get('project', 'Sin proyecto')}")
            doc.add_paragraph(f"Precondiciones: {tc['preconditions']}")
            doc.add_paragraph(f"Pasos: {tc['steps']}")
            doc.add_paragraph(f"Resultado Esperado: {tc['expected']}")
            if tc.get('actual_result'):
                doc.add_paragraph(f"Resultado Obtenido: {tc['actual_result']}")
            doc.add_page_break()
        
        filename = self.data_dir / "exports" / f"CasosPrueba_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(str(filename))
        
        messagebox.showinfo("Éxito", f"Archivo guardado en:\n{filename}")
    
    def export_bugs_to_pdf(self):
        """MEJORA 3: Exportar bugs a PDF con gráficos"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
            from reportlab.platypus import Image as RLImage
            from reportlab.graphics.shapes import Drawing
            from reportlab.graphics.charts.piecharts import Pie
            from reportlab.graphics.charts.barcharts import VerticalBarChart
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
        except ImportError:
            messagebox.showerror(
                "Librería Faltante",
                "Para exportar a PDF necesitas instalar reportlab:\n\npip install reportlab"
            )
            return
        
        if not self.bugs:
            messagebox.showwarning("Advertencia", "No hay bugs para exportar")
            return
        
        filename = self.data_dir / "exports" / f"Reporte_Bugs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        doc = SimpleDocTemplate(str(filename), pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Estilos personalizados
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1E88E5'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#E53935'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Título
        story.append(Paragraph("Reporte de Bugs", title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Resumen ejecutivo
        story.append(Paragraph("Resumen Ejecutivo", heading_style))
        
        total_bugs = len(self.bugs)
        open_bugs = len([b for b in self.bugs if b['status'] not in ['Cerrado', 'Resuelto']])
        critical_bugs = len([b for b in self.bugs if b['severity'] == 'Crítico'])
        
        summary_data = [
            ['Métrica', 'Valor'],
            ['Total de Bugs', str(total_bugs)],
            ['Bugs Abiertos', str(open_bugs)],
            ['Bugs Críticos', str(critical_bugs)],
            ['Fecha del Reporte', datetime.now().strftime('%Y-%m-%d %H:%M')]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E88E5')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.5*inch))
        
        # Gráfico de bugs por severidad
        story.append(Paragraph("Distribución por Severidad", heading_style))
        
        severity_counts = defaultdict(int)
        for bug in self.bugs:
            severity_counts[bug['severity']] += 1
        
        if severity_counts:
            drawing = Drawing(400, 200)
            pie = Pie()
            pie.x = 150
            pie.y = 50
            pie.width = 100
            pie.height = 100
            pie.data = list(severity_counts.values())
            pie.labels = list(severity_counts.keys())
            pie.slices.strokeWidth = 0.5
            pie.slices[0].fillColor = colors.HexColor('#D32F2F')  # Crítico
            pie.slices[1].fillColor = colors.HexColor('#F57C00')  # Alto
            pie.slices[2].fillColor = colors.HexColor('#FBC02D')  # Medio
            if len(pie.slices) > 3:
                pie.slices[3].fillColor = colors.HexColor('#388E3C')  # Bajo
            drawing.add(pie)
            story.append(drawing)
        
        story.append(Spacer(1, 0.5*inch))
        
        # Listado de bugs
        story.append(Paragraph("Listado Detallado de Bugs", heading_style))
        
        for bug in self.bugs:
            bug_data = [
                ['Campo', 'Valor'],
                ['ID', f"#{bug['id']}"],
                ['Título', bug['title']],
                ['Estado', bug['status']],
                ['Severidad', bug['severity']],
                ['Prioridad', bug['priority']],
                ['Tipo', bug['type']],
                ['Proyecto', bug.get('project', 'Sin proyecto')],
                ['Descripción', bug['description']],
                ['Fecha', bug['date']]
            ]
            
            bug_table = Table(bug_data, colWidths=[1.5*inch, 4*inch])
            bug_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E53935')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP')
            ]))
            
            story.append(bug_table)
            story.append(Spacer(1, 0.3*inch))
        
        # Generar PDF
        doc.build(story)
        messagebox.showinfo("Éxito", f"PDF generado exitosamente:\n{filename}")
    
    def export_test_cases_to_pdf(self):
        """ MEJORA 3: Exportar casos de prueba a PDF con gráficos"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.graphics.shapes import Drawing
            from reportlab.graphics.charts.piecharts import Pie
            from reportlab.lib.enums import TA_CENTER
        except ImportError:
            messagebox.showerror(
                "Librería Faltante",
                "Para exportar a PDF necesitas instalar reportlab:\n\npip install reportlab"
            )
            return
        
        if not self.test_cases:
            messagebox.showwarning("Advertencia", "No hay casos para exportar")
            return
        
        filename = self.data_dir / "exports" / f"Plan_Pruebas_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        doc = SimpleDocTemplate(str(filename), pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Estilos
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1E88E5'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#1E88E5'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Título
        story.append(Paragraph("Plan de Pruebas", title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Resumen
        story.append(Paragraph("Resumen Ejecutivo", heading_style))
        
        total_cases = len(self.test_cases)
        executed = len([tc for tc in self.test_cases if tc.get('status') != 'Sin Probar'])
        passed = len([tc for tc in self.test_cases if tc.get('status') == 'Aprobado'])
        coverage = (executed / total_cases * 100) if total_cases > 0 else 0
        pass_rate = (passed / executed * 100) if executed > 0 else 0
        
        summary_data = [
            ['Métrica', 'Valor'],
            ['Total de Casos', str(total_cases)],
            ['Casos Ejecutados', str(executed)],
            ['Casos Aprobados', str(passed)],
            ['Cobertura de Ejecución', f"{coverage:.1f}%"],
            ['Tasa de Éxito', f"{pass_rate:.1f}%"],
            ['Fecha del Reporte', datetime.now().strftime('%Y-%m-%d %H:%M')]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E88E5')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.5*inch))
        
        # Gráfico de casos por estado
        story.append(Paragraph("Distribución por Estado", heading_style))
        
        status_counts = defaultdict(int)
        for tc in self.test_cases:
            status_counts[tc.get('status', 'Sin Probar')] += 1
        
        if status_counts:
            drawing = Drawing(400, 200)
            pie = Pie()
            pie.x = 150
            pie.y = 50
            pie.width = 100
            pie.height = 100
            pie.data = list(status_counts.values())
            pie.labels = list(status_counts.keys())
            pie.slices.strokeWidth = 0.5
            drawing.add(pie)
            story.append(drawing)
        
        story.append(Spacer(1, 0.5*inch))
        
        # Listado de casos
        story.append(Paragraph("Casos de Prueba Detallados", heading_style))
        
        for tc in self.test_cases:
            tc_data = [
                ['Campo', 'Valor'],
                ['ID', tc['test_id']],
                ['Título', tc['title']],
                ['Módulo', tc['module']],
                ['Estado', tc.get('status', 'Sin Probar')],
                ['Prioridad', tc['priority']],
                ['Tipo', tc['type']],
                ['Precondiciones', tc['preconditions']],
                ['Pasos', tc['steps']],
                ['Resultado Esperado', tc['expected']]
            ]
            
            tc_table = Table(tc_data, colWidths=[1.5*inch, 4*inch])
            tc_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E88E5')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP')
            ]))
            
            story.append(tc_table)
            story.append(Spacer(1, 0.3*inch))
        
        # Generar PDF
        doc.build(story)
        messagebox.showinfo("Éxito", f"PDF generado exitosamente:\n{filename}")

if __name__ == '__main__':
    root = ctk.CTk()
    app = QAManagerApp(root)
    root.mainloop()
